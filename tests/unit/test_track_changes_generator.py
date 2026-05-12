"""Tests unitaires pour le TrackChangesGenerator.

Teste la génération des marques de révision OOXML (w:ins / w:del) pour
les cas simples de remplacement, insertion et suppression de mots, ainsi
que la structure XML produite et la préservation du formatage.

Valide : Exigences 5.2, 5.4, 5.5
"""

import io
import sys
import zipfile
from pathlib import Path

import pytest
from lxml import etree

# Ajouter le backend au path
sys.path.insert(
    0, str(Path(__file__).resolve().parents[2] / "local-site" / "web" / "backend")
)

from services.document_parser import DocumentParser
from services.revision_models import ParagraphCorrection, ParsedDocument
from services.track_changes_generator import (
    NSMAP,
    WORD_NS,
    TrackChangesGenerator,
    W_DEL,
    W_DEL_TEXT,
    W_INS,
    W_R,
    W_RPR,
    W_T,
)

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

XML_NS = "http://www.w3.org/XML/1998/namespace"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _create_simple_docx(paragraphs: list[str]) -> bytes:
    """Crée un fichier .docx minimal avec les paragraphes donnés.

    Args:
        paragraphs: Liste de textes pour chaque paragraphe.

    Returns:
        Bytes du fichier .docx.
    """
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _create_formatted_docx(paragraphs: list[dict]) -> bytes:
    """Crée un .docx avec des runs formatés.

    Args:
        paragraphs: Liste de dicts avec 'runs' contenant
                    [{'text': str, 'bold': bool, 'italic': bool}].

    Returns:
        Bytes du fichier .docx.
    """
    from docx import Document

    doc = Document()
    for para_def in paragraphs:
        para = doc.add_paragraph()
        for run_def in para_def["runs"]:
            run = para.add_run(run_def["text"])
            if run_def.get("bold"):
                run.bold = True
            if run_def.get("italic"):
                run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _parse_output_xml(output_bytes: bytes) -> etree._Element:
    """Extrait et parse le document.xml du .docx de sortie.

    Args:
        output_bytes: Bytes du fichier .docx produit.

    Returns:
        Élément racine du document XML.
    """
    with zipfile.ZipFile(io.BytesIO(output_bytes), "r") as zf:
        xml_bytes = zf.read("word/document.xml")
    return etree.fromstring(xml_bytes)


def _find_all_elements(root: etree._Element, tag: str) -> list[etree._Element]:
    """Trouve tous les éléments avec le tag donné (recherche récursive).

    Args:
        root: Élément racine XML.
        tag: Tag qualifié à rechercher.

    Returns:
        Liste des éléments trouvés.
    """
    return root.findall(f".//{tag}")


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def generator() -> TrackChangesGenerator:
    """Instance fraîche du TrackChangesGenerator."""
    return TrackChangesGenerator()


@pytest.fixture
def parser() -> DocumentParser:
    """Instance du DocumentParser."""
    return DocumentParser()


# ---------------------------------------------------------------------------
# Tests : Remplacement simple de mot
# ---------------------------------------------------------------------------


class TestSimpleWordReplacement:
    """Teste le remplacement d'un mot (w:del + w:ins)."""

    def test_replacement_generates_del_and_ins(self, generator, parser):
        """Un remplacement de mot produit un w:del et un w:ins."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        ins_elements = _find_all_elements(root, W_INS)

        assert len(del_elements) >= 1, "Au moins un w:del attendu"
        assert len(ins_elements) >= 1, "Au moins un w:ins attendu"

    def test_replacement_del_contains_original_word(self, generator, parser):
        """Le w:del contient le mot original supprimé."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        del_texts = []
        for del_elem in del_elements:
            for dt in del_elem.findall(f".//{W_DEL_TEXT}"):
                if dt.text:
                    del_texts.append(dt.text)

        combined_del = "".join(del_texts)
        assert "noir" in combined_del

    def test_replacement_ins_contains_new_word(self, generator, parser):
        """Le w:ins contient le mot de remplacement."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        ins_elements = _find_all_elements(root, W_INS)
        ins_texts = []
        for ins_elem in ins_elements:
            for t in ins_elem.findall(f".//{W_T}"):
                if t.text:
                    ins_texts.append(t.text)

        combined_ins = "".join(ins_texts)
        assert "blanc" in combined_ins


# ---------------------------------------------------------------------------
# Tests : Insertion de mot
# ---------------------------------------------------------------------------


class TestWordInsertion:
    """Teste l'insertion d'un mot (w:ins sans w:del correspondant)."""

    def test_insertion_generates_ins_element(self, generator, parser):
        """L'insertion d'un mot produit un w:ins."""
        docx_bytes = _create_simple_docx(["Le chat"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat",
                corrected_text="Le gros chat",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        ins_elements = _find_all_elements(root, W_INS)
        assert len(ins_elements) >= 1

        ins_texts = []
        for ins_elem in ins_elements:
            for t in ins_elem.findall(f".//{W_T}"):
                if t.text:
                    ins_texts.append(t.text)

        combined_ins = "".join(ins_texts)
        assert "gros" in combined_ins


# ---------------------------------------------------------------------------
# Tests : Suppression de mot
# ---------------------------------------------------------------------------


class TestWordDeletion:
    """Teste la suppression d'un mot (w:del sans w:ins correspondant)."""

    def test_deletion_generates_del_element(self, generator, parser):
        """La suppression d'un mot produit un w:del."""
        docx_bytes = _create_simple_docx(["Le gros chat"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le gros chat",
                corrected_text="Le chat",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        assert len(del_elements) >= 1

        del_texts = []
        for del_elem in del_elements:
            for dt in del_elem.findall(f".//{W_DEL_TEXT}"):
                if dt.text:
                    del_texts.append(dt.text)

        combined_del = "".join(del_texts)
        assert "gros" in combined_del


# ---------------------------------------------------------------------------
# Tests : Paragraphe inchangé
# ---------------------------------------------------------------------------


class TestUnchangedParagraph:
    """Teste qu'un paragraphe sans correction n'est pas modifié."""

    def test_unchanged_paragraph_has_no_revision_marks(self, generator, parser):
        """Un paragraphe sans correction ne contient ni w:del ni w:ins."""
        docx_bytes = _create_simple_docx(["Texte identique", "Autre paragraphe"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        # Seul le second paragraphe est corrigé
        corrections = [
            ParagraphCorrection(
                paragraph_index=1,
                original_text="Autre paragraphe",
                corrected_text="Autre section",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        # Trouver le premier paragraphe (après le body)
        body = root.find(f"{{{WORD_NS}}}body")
        paragraphs = body.findall(f"{{{WORD_NS}}}p")

        # Le premier paragraphe ne doit pas avoir de w:del ou w:ins
        first_para = paragraphs[0]
        del_in_first = first_para.findall(f".//{W_DEL}")
        ins_in_first = first_para.findall(f".//{W_INS}")

        assert len(del_in_first) == 0, "Pas de w:del dans le paragraphe inchangé"
        assert len(ins_in_first) == 0, "Pas de w:ins dans le paragraphe inchangé"

    def test_no_corrections_produces_unchanged_output(self, generator, parser):
        """Sans aucune correction, le document ne contient aucune marque."""
        docx_bytes = _create_simple_docx(["Premier paragraphe", "Deuxième paragraphe"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = []  # Aucune correction

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        ins_elements = _find_all_elements(root, W_INS)

        assert len(del_elements) == 0
        assert len(ins_elements) == 0


# ---------------------------------------------------------------------------
# Tests : Structure XML des éléments de révision
# ---------------------------------------------------------------------------


class TestXMLStructure:
    """Teste la structure XML des éléments w:del et w:ins."""

    def test_del_has_required_attributes(self, generator, parser):
        """w:del a les attributs w:id, w:author et w:date."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        assert len(del_elements) >= 1

        for del_elem in del_elements:
            assert del_elem.get(f"{{{WORD_NS}}}id") is not None, "w:id manquant"
            assert del_elem.get(f"{{{WORD_NS}}}author") is not None, "w:author manquant"
            assert del_elem.get(f"{{{WORD_NS}}}date") is not None, "w:date manquant"

    def test_ins_has_required_attributes(self, generator, parser):
        """w:ins a les attributs w:id, w:author et w:date."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        ins_elements = _find_all_elements(root, W_INS)
        assert len(ins_elements) >= 1

        for ins_elem in ins_elements:
            assert ins_elem.get(f"{{{WORD_NS}}}id") is not None, "w:id manquant"
            assert ins_elem.get(f"{{{WORD_NS}}}author") is not None, "w:author manquant"
            assert ins_elem.get(f"{{{WORD_NS}}}date") is not None, "w:date manquant"

    def test_author_is_judi_expert(self, generator, parser):
        """L'auteur des révisions est 'Judi-Expert'."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        all_revisions = _find_all_elements(root, W_DEL) + _find_all_elements(root, W_INS)
        for rev in all_revisions:
            assert rev.get(f"{{{WORD_NS}}}author") == "Judi-Expert"

    def test_revision_ids_are_unique(self, generator, parser):
        """Tous les w:id sont uniques dans le document."""
        docx_bytes = _create_simple_docx(["Le chat noir", "Un gros chien"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            ),
            ParagraphCorrection(
                paragraph_index=1,
                original_text="Un gros chien",
                corrected_text="Un petit chien",
                has_changes=True,
            ),
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        all_revisions = _find_all_elements(root, W_DEL) + _find_all_elements(root, W_INS)
        ids = [rev.get(f"{{{WORD_NS}}}id") for rev in all_revisions]

        assert len(ids) == len(set(ids)), f"IDs non uniques : {ids}"

    def test_del_contains_del_text_element(self, generator, parser):
        """w:del contient un w:r avec un w:delText."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        for del_elem in del_elements:
            runs = del_elem.findall(f".//{W_R}")
            assert len(runs) >= 1, "w:del doit contenir au moins un w:r"
            del_texts = del_elem.findall(f".//{W_DEL_TEXT}")
            assert len(del_texts) >= 1, "w:del/w:r doit contenir un w:delText"

    def test_ins_contains_t_element(self, generator, parser):
        """w:ins contient un w:r avec un w:t."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        ins_elements = _find_all_elements(root, W_INS)
        for ins_elem in ins_elements:
            runs = ins_elem.findall(f".//{W_R}")
            assert len(runs) >= 1, "w:ins doit contenir au moins un w:r"
            t_elements = ins_elem.findall(f".//{W_T}")
            assert len(t_elements) >= 1, "w:ins/w:r doit contenir un w:t"

    def test_date_is_iso_format(self, generator, parser):
        """L'attribut w:date est au format ISO 8601."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        all_revisions = _find_all_elements(root, W_DEL) + _find_all_elements(root, W_INS)
        for rev in all_revisions:
            date_val = rev.get(f"{{{WORD_NS}}}date")
            assert date_val is not None
            # Format attendu : 2024-01-15T10:30:00Z
            assert date_val.endswith("Z"), f"Date doit finir par Z : {date_val}"
            assert "T" in date_val, f"Date doit contenir T : {date_val}"


# ---------------------------------------------------------------------------
# Tests : Corrections multiples sur plusieurs paragraphes
# ---------------------------------------------------------------------------


class TestMultipleCorrections:
    """Teste les corrections sur plusieurs paragraphes."""

    def test_multiple_paragraphs_corrected(self, generator, parser):
        """Plusieurs paragraphes peuvent être corrigés indépendamment."""
        docx_bytes = _create_simple_docx([
            "Le chat noir",
            "Un gros chien",
            "La petite souris",
        ])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            ),
            ParagraphCorrection(
                paragraph_index=2,
                original_text="La petite souris",
                corrected_text="La grande souris",
                has_changes=True,
            ),
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        body = root.find(f"{{{WORD_NS}}}body")
        paragraphs = body.findall(f"{{{WORD_NS}}}p")

        # Paragraphe 0 : corrigé (a des marques)
        assert len(paragraphs[0].findall(f".//{W_DEL}")) >= 1
        assert len(paragraphs[0].findall(f".//{W_INS}")) >= 1

        # Paragraphe 1 : inchangé (pas de marques)
        assert len(paragraphs[1].findall(f".//{W_DEL}")) == 0
        assert len(paragraphs[1].findall(f".//{W_INS}")) == 0

        # Paragraphe 2 : corrigé (a des marques)
        assert len(paragraphs[2].findall(f".//{W_DEL}")) >= 1
        assert len(paragraphs[2].findall(f".//{W_INS}")) >= 1

    def test_has_changes_false_is_skipped(self, generator, parser):
        """Les corrections avec has_changes=False sont ignorées."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat noir",
                has_changes=False,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        ins_elements = _find_all_elements(root, W_INS)

        assert len(del_elements) == 0
        assert len(ins_elements) == 0


# ---------------------------------------------------------------------------
# Tests : Préservation du formatage
# ---------------------------------------------------------------------------


class TestFormattingPreservation:
    """Teste que le formatage (w:rPr) est préservé sur les révisions."""

    def test_del_element_has_rpr(self, generator, parser):
        """w:del préserve les propriétés de formatage (w:rPr)."""
        docx_bytes = _create_formatted_docx([
            {"runs": [{"text": "Le chat noir", "bold": True, "italic": False}]}
        ])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        del_elements = _find_all_elements(root, W_DEL)
        assert len(del_elements) >= 1

        # Vérifier que le run dans w:del a un w:rPr
        for del_elem in del_elements:
            runs = del_elem.findall(f".//{W_R}")
            for run in runs:
                rpr = run.find(W_RPR)
                assert rpr is not None, "w:rPr manquant dans w:del/w:r"

    def test_ins_element_has_rpr(self, generator, parser):
        """w:ins préserve les propriétés de formatage (w:rPr)."""
        docx_bytes = _create_formatted_docx([
            {"runs": [{"text": "Le chat noir", "bold": True, "italic": False}]}
        ])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        ins_elements = _find_all_elements(root, W_INS)
        assert len(ins_elements) >= 1

        # Vérifier que le run dans w:ins a un w:rPr
        for ins_elem in ins_elements:
            runs = ins_elem.findall(f".//{W_R}")
            for run in runs:
                rpr = run.find(W_RPR)
                assert rpr is not None, "w:rPr manquant dans w:ins/w:r"

    def test_formatting_preserved_on_equal_runs(self, generator, parser):
        """Les runs inchangés conservent aussi le formatage."""
        docx_bytes = _create_formatted_docx([
            {"runs": [{"text": "Le chat noir", "bold": True, "italic": True}]}
        ])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        root = _parse_output_xml(output)

        body = root.find(f"{{{WORD_NS}}}body")
        para = body.findall(f"{{{WORD_NS}}}p")[0]

        # Les runs normaux (texte inchangé "Le chat ") doivent avoir w:rPr
        normal_runs = para.findall(W_R)
        if normal_runs:
            for run in normal_runs:
                rpr = run.find(W_RPR)
                assert rpr is not None, "w:rPr manquant sur run inchangé"


# ---------------------------------------------------------------------------
# Tests : Sortie est un .docx valide
# ---------------------------------------------------------------------------


class TestOutputValidity:
    """Teste que la sortie est un fichier .docx valide."""

    def test_output_is_valid_zip(self, generator, parser):
        """La sortie est une archive ZIP valide."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        assert zipfile.is_zipfile(io.BytesIO(output))

    def test_output_contains_document_xml(self, generator, parser):
        """La sortie contient word/document.xml."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)
        with zipfile.ZipFile(io.BytesIO(output), "r") as zf:
            assert "word/document.xml" in zf.namelist()

    def test_output_can_be_reparsed(self, generator, parser):
        """Le .docx de sortie peut être re-parsé par DocumentParser."""
        docx_bytes = _create_simple_docx(["Le chat noir"])
        parsed_doc = parser.parse(docx_bytes, "docx")

        corrections = [
            ParagraphCorrection(
                paragraph_index=0,
                original_text="Le chat noir",
                corrected_text="Le chat blanc",
                has_changes=True,
            )
        ]

        output = generator.generate(parsed_doc, corrections)

        # Le fichier de sortie doit pouvoir être re-parsé sans erreur
        reparsed = parser.parse(output, "docx")
        assert isinstance(reparsed, ParsedDocument)
