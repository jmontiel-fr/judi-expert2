"""Tests unitaires pour le TREParser.

Teste le parsing de documents TRE (.docx) : extraction de placeholders,
annotations, validation de structure, extraction PE et en-tête.
"""

import sys
from pathlib import Path

import pytest
from docx import Document as DocxDocument

# Ajouter le backend au path
sys.path.insert(
    0, str(Path(__file__).resolve().parents[2] / "local-site" / "web" / "backend")
)

from services.tre_parser import Annotation, Placeholder, TREParser, TREParseResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _create_docx(tmp_path: Path, paragraphs: list[str], filename: str = "test.docx") -> str:
    """Crée un fichier .docx avec les paragraphes donnés.

    Args:
        tmp_path: Répertoire temporaire pytest.
        paragraphs: Liste de textes de paragraphes.
        filename: Nom du fichier à créer.

    Returns:
        Chemin absolu vers le fichier créé.
    """
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    path = tmp_path / filename
    doc.save(str(path))
    return str(path)


# ---------------------------------------------------------------------------
# Tests : Parsing valide
# ---------------------------------------------------------------------------


class TestParseValidTRE:
    """Vérifie le parsing d'un TRE valide avec placeholders et annotations."""

    def test_parse_valid_tre(self, tmp_path):
        """Un TRE avec placeholders et annotations est correctement parsé."""
        paragraphs = [
            "En-tête du rapport",
            "Expert : <<nom_expert>>",
            "Dossier : <<numero_dossier>>",
            "@debut_tpe@",
            "@dires Contenu des dires de la partie@",
            "@analyse Analyse du contenu@",
            "Paragraphe normal sans annotation.",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)

        # Vérifier les placeholders
        assert len(result.placeholders) == 2
        assert result.placeholders[0].name == "nom_expert"
        assert result.placeholders[1].name == "numero_dossier"

        # Vérifier les annotations (debut_tpe + dires + analyse)
        assert len(result.annotations) == 3
        assert result.annotations[0].type == "debut_tpe"
        assert result.annotations[1].type == "dires"
        assert result.annotations[1].content == "Contenu des dires de la partie"
        assert result.annotations[2].type == "analyse"
        assert result.annotations[2].content == "Analyse du contenu"

        # Vérifier la position de @debut_tpe@
        assert result.debut_tpe_position == 3

        # Pas d'erreurs
        assert result.errors == []


# ---------------------------------------------------------------------------
# Tests : Validation — @debut_tpe@ manquant
# ---------------------------------------------------------------------------


class TestValidateMissingDebutTpe:
    """Vérifie que l'absence de @debut_tpe@ est signalée."""

    def test_validate_missing_debut_tpe(self, tmp_path):
        """Un TRE sans @debut_tpe@ produit une erreur de validation."""
        paragraphs = [
            "En-tête du rapport",
            "Expert : <<nom_expert>>",
            "@dires Contenu des dires@",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)
        errors = parser.validate(result)

        assert any("@debut_tpe@" in e and "absent" in e for e in errors)


# ---------------------------------------------------------------------------
# Tests : Validation — Annotation non fermée
# ---------------------------------------------------------------------------


class TestValidateUnclosedAnnotation:
    """Vérifie qu'une annotation non fermée est signalée."""

    def test_validate_unclosed_annotation(self, tmp_path):
        """Une annotation ouverte sans @ de fermeture produit une erreur."""
        paragraphs = [
            "@debut_tpe@",
            "@dires Début du contenu sans fermeture",
            "Suite du contenu toujours sans fermeture",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)
        errors = parser.validate(result)

        assert any("non fermée" in e for e in errors)


# ---------------------------------------------------------------------------
# Tests : Extraction PE avec questions
# ---------------------------------------------------------------------------


class TestExtractPEWithQuestions:
    """Vérifie l'extraction du PE avec ajout des questions en conclusion."""

    def test_extract_pe_with_questions(self, tmp_path):
        """Le PE contient les paragraphes après @debut_tpe@ et les questions."""
        paragraphs = [
            "En-tête à exclure",
            "Autre en-tête",
            "@debut_tpe@",
            "Premier paragraphe du TPE",
            "Deuxième paragraphe du TPE",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        questions = {
            "question_1": "Quel est le préjudice subi ?",
            "question_2": "Quelles sont les séquelles ?",
        }

        pe_bytes = parser.extract_pe(docx_path, questions)

        # Vérifier que le résultat est un docx valide
        import io

        pe_doc = DocxDocument(io.BytesIO(pe_bytes))
        texts = [p.text for p in pe_doc.paragraphs if p.text.strip()]

        # Les paragraphes après @debut_tpe@ sont présents
        assert "Premier paragraphe du TPE" in texts
        assert "Deuxième paragraphe du TPE" in texts

        # La section Conclusion est présente
        assert "Conclusion" in texts

        # Les questions sont présentes
        assert "Quel est le préjudice subi ?" in texts
        assert "Quelles sont les séquelles ?" in texts

        # L'en-tête n'est PAS présent
        assert "En-tête à exclure" not in texts


# ---------------------------------------------------------------------------
# Tests : Extraction en-tête
# ---------------------------------------------------------------------------


class TestExtractHeader:
    """Vérifie l'extraction de l'en-tête (avant @debut_tpe@)."""

    def test_extract_header(self, tmp_path):
        """L'en-tête contient uniquement les paragraphes avant @debut_tpe@."""
        paragraphs = [
            "Tribunal de Grande Instance",
            "Expert : Dr. Dupont",
            "Référence : 2024/12345",
            "@debut_tpe@",
            "Contenu du TPE à exclure",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        header_bytes = parser.extract_header(docx_path)

        import io

        header_doc = DocxDocument(io.BytesIO(header_bytes))
        texts = [p.text for p in header_doc.paragraphs if p.text.strip()]

        # Les paragraphes d'en-tête sont présents
        assert "Tribunal de Grande Instance" in texts
        assert "Expert : Dr. Dupont" in texts
        assert "Référence : 2024/12345" in texts

        # Le contenu après @debut_tpe@ n'est PAS présent
        assert "Contenu du TPE à exclure" not in texts
        assert "@debut_tpe@" not in texts


# ---------------------------------------------------------------------------
# Tests : Annotation personnalisée (@/custom)
# ---------------------------------------------------------------------------


class TestParseCustomAnnotation:
    """Vérifie le parsing des annotations personnalisées."""

    def test_parse_custom_annotation(self, tmp_path):
        """Une annotation @/custom est parsée avec is_custom=True."""
        paragraphs = [
            "@debut_tpe@",
            "@/mon_annotation Contenu personnalisé@",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)

        # Trouver l'annotation personnalisée (pas debut_tpe)
        custom_annotations = [a for a in result.annotations if a.is_custom]
        assert len(custom_annotations) == 1
        assert custom_annotations[0].type == "mon_annotation"
        assert custom_annotations[0].content == "Contenu personnalisé"
        assert custom_annotations[0].is_custom is True


# ---------------------------------------------------------------------------
# Tests : Annotation multi-paragraphe
# ---------------------------------------------------------------------------


class TestParseMultilineAnnotation:
    """Vérifie le parsing des annotations sur plusieurs paragraphes."""

    def test_parse_multiline_annotation(self, tmp_path):
        """Une annotation multi-paragraphe est correctement assemblée."""
        paragraphs = [
            "@debut_tpe@",
            "@dires Première ligne du contenu",
            "Deuxième ligne du contenu",
            "Troisième ligne avec fermeture @",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)

        # Trouver l'annotation dires
        dires_annotations = [a for a in result.annotations if a.type == "dires"]
        assert len(dires_annotations) == 1

        content = dires_annotations[0].content
        assert "Première ligne du contenu" in content
        assert "Deuxième ligne du contenu" in content
        assert "Troisième ligne avec fermeture" in content


# ---------------------------------------------------------------------------
# Tests : Validation snake_case des placeholders
# ---------------------------------------------------------------------------


class TestPlaceholderSnakeCaseValidation:
    """Vérifie que les placeholders non-snake_case sont signalés."""

    def test_placeholder_snake_case_validation(self, tmp_path):
        """Un placeholder non-snake_case produit une erreur de validation."""
        paragraphs = [
            "@debut_tpe@",
            "Valide : <<nom_expert>>",
            "Invalide : <<NomExpert>>",
            "Invalide : <<nomExpert>>",
        ]
        docx_path = _create_docx(tmp_path, paragraphs)
        parser = TREParser()

        result = parser.parse(docx_path)
        errors = parser.validate(result)

        # Le placeholder valide ne génère pas d'erreur
        assert not any("nom_expert" in e for e in errors)

        # Les placeholders invalides génèrent des erreurs
        assert any("NomExpert" in e and "snake_case" in e for e in errors)
        assert any("nomExpert" in e and "snake_case" in e for e in errors)
