"""Tests par propriété — Track Changes Generator (OOXML revision marks).

Propriétés testées :
- Property 2 : Préservation du formatage des runs
- Property 4 : Round-trip du contenu textuel

Utilise Hypothesis pour générer des documents .docx avec formatage varié
et des paires (texte original, texte corrigé), puis vérifie que :
- Le formatage est préservé sur les éléments w:del et w:ins (Property 2)
- L'extraction du texte accepté/rejeté depuis le XML produit correspond
  respectivement au texte corrigé et au texte original (Property 4)
"""

from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path
from typing import Any

from hypothesis import assume, given, settings
from hypothesis import strategies as st
from lxml import etree

# Ajouter le backend au path
sys.path.insert(
    0,
    str(
        Path(__file__).resolve().parents[2]
        / "local-site"
        / "web"
        / "backend"
    ),
)

from services.document_parser import DocumentParser
from services.revision_models import ParagraphCorrection, ParsedDocument
from services.track_changes_generator import TrackChangesGenerator

# ---------------------------------------------------------------------------
# Constantes OOXML
# ---------------------------------------------------------------------------

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


def _qn(tag: str) -> str:
    """Construit un nom qualifié Clark notation pour le namespace w:."""
    return f"{{{WORD_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Helpers d'extraction de texte depuis le XML de sortie
# ---------------------------------------------------------------------------


def extract_accepted_text_from_paragraph(p_element: "etree._Element") -> str:
    """Extrait le texte "accepté" d'un paragraphe avec track changes.

    Le texte accepté correspond à ce qu'on obtient en acceptant toutes les
    révisions : on garde les insertions (w:ins) et le texte normal (w:r),
    on supprime les suppressions (w:del).

    Args:
        p_element: Élément XML w:p.

    Returns:
        Texte accepté du paragraphe.
    """
    parts: list[str] = []

    for child in p_element:
        tag = etree.QName(child.tag).localname

        if tag == "del":
            # Suppression → on ignore (accepter = retirer le texte supprimé)
            continue
        elif tag == "ins":
            # Insertion → on garde le texte (accepter = garder l'insertion)
            for r_elem in child.findall(_qn("r")):
                for t_elem in r_elem.findall(_qn("t")):
                    if t_elem.text:
                        parts.append(t_elem.text)
        elif tag == "r":
            # Run normal (texte inchangé) → on garde
            for t_elem in child.findall(_qn("t")):
                if t_elem.text:
                    parts.append(t_elem.text)
        # Ignorer w:pPr et autres éléments non-textuels

    return "".join(parts)


def extract_rejected_text_from_paragraph(p_element: "etree._Element") -> str:
    """Extrait le texte "rejeté" d'un paragraphe avec track changes.

    Le texte rejeté correspond à ce qu'on obtient en rejetant toutes les
    révisions : on garde les suppressions (w:del) et le texte normal (w:r),
    on supprime les insertions (w:ins).

    Args:
        p_element: Élément XML w:p.

    Returns:
        Texte rejeté du paragraphe.
    """
    parts: list[str] = []

    for child in p_element:
        tag = etree.QName(child.tag).localname

        if tag == "ins":
            # Insertion → on ignore (rejeter = retirer l'insertion)
            continue
        elif tag == "del":
            # Suppression → on garde le texte (rejeter = garder le texte supprimé)
            for r_elem in child.findall(_qn("r")):
                for dt_elem in r_elem.findall(_qn("delText")):
                    if dt_elem.text:
                        parts.append(dt_elem.text)
        elif tag == "r":
            # Run normal (texte inchangé) → on garde
            for t_elem in child.findall(_qn("t")):
                if t_elem.text:
                    parts.append(t_elem.text)
        # Ignorer w:pPr et autres éléments non-textuels

    return "".join(parts)


def _find_paragraph_with_revisions(
    xml_bytes: bytes,
) -> "etree._Element | None":
    """Trouve le premier paragraphe contenant des éléments de révision.

    Args:
        xml_bytes: Bytes du fichier .docx.

    Returns:
        Élément w:p contenant des révisions, ou None.
    """
    with zipfile.ZipFile(io.BytesIO(xml_bytes), "r") as zf:
        doc_xml = zf.read("word/document.xml")

    root = etree.fromstring(doc_xml)
    all_paragraphs = root.findall(f".//{{{WORD_NS}}}p")

    for para in all_paragraphs:
        del_elems = para.findall(f"{{{WORD_NS}}}del")
        ins_elems = para.findall(f"{{{WORD_NS}}}ins")
        if del_elems or ins_elems:
            return para

    return None


# ---------------------------------------------------------------------------
# Hypothesis strategies — Property 4 (round-trip)
# ---------------------------------------------------------------------------

# Alphabet pour le texte : lettres, chiffres, ponctuation, espaces
_text_alphabet = st.characters(
    whitelist_categories=("L", "N", "P", "Zs"),
    blacklist_characters="\x00",
)


@st.composite
def paragraph_correction_pairs(
    draw: st.DrawFn,
) -> tuple[bytes, str, str]:
    """Génère un document .docx avec un paragraphe et un texte corrigé différent.

    Returns:
        Tuple (docx_bytes, original_text, corrected_text).
    """
    from docx import Document

    # Générer le texte original (au moins 2 caractères pour avoir du contenu)
    original_text = draw(
        st.text(alphabet=_text_alphabet, min_size=2, max_size=150)
    ).strip()
    assume(len(original_text) >= 2)

    # Générer le texte corrigé (différent de l'original)
    corrected_text = draw(
        st.text(alphabet=_text_alphabet, min_size=2, max_size=150)
    ).strip()
    assume(len(corrected_text) >= 2)
    assume(original_text != corrected_text)

    # Créer un document .docx avec un seul paragraphe contenant le texte original
    doc = Document()
    doc.add_paragraph(original_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    return docx_bytes, original_text, corrected_text


# ---------------------------------------------------------------------------
# Property 4 : Round-trip du contenu textuel
# ---------------------------------------------------------------------------


@settings(max_examples=100, deadline=None)
@given(data=paragraph_correction_pairs())
def test_textual_content_round_trip(
    data: tuple[bytes, str, str],
) -> None:
    """Pour tout document, l'extraction du texte accepté depuis les track changes
    DOIT être égale au texte corrigé, et l'extraction du texte rejeté DOIT être
    égale au texte original.

    **Validates: Requirements 5.2, 5.5**
    """
    docx_bytes, original_text, corrected_text = data

    # 1. Parser le document
    parser = DocumentParser()
    parsed_doc = parser.parse(docx_bytes, "docx")

    # 2. Trouver le paragraphe contenant le texte original
    target_para = None
    for para in parsed_doc.paragraphs:
        if para.full_text == original_text:
            target_para = para
            break

    assert target_para is not None, (
        f"Could not find paragraph with text {original_text!r} in parsed doc"
    )

    paragraph_index = target_para.index

    # 3. Créer la correction
    correction = ParagraphCorrection(
        paragraph_index=paragraph_index,
        original_text=original_text,
        corrected_text=corrected_text,
        has_changes=True,
    )

    # 4. Générer les track changes
    generator = TrackChangesGenerator()
    output_bytes = generator.generate(parsed_doc, [correction])

    # 5. Trouver le paragraphe modifié dans le XML de sortie
    target_paragraph = _find_paragraph_with_revisions(output_bytes)
    assert target_paragraph is not None, (
        "No paragraph with revision elements found in output"
    )

    # 6. Extraire le texte accepté et rejeté
    accepted_text = extract_accepted_text_from_paragraph(target_paragraph)
    rejected_text = extract_rejected_text_from_paragraph(target_paragraph)

    # 7. Vérifier le round-trip
    assert accepted_text == corrected_text, (
        f"Accepted text mismatch.\n"
        f"  Expected (corrected): {corrected_text!r}\n"
        f"  Got (accepted):       {accepted_text!r}"
    )
    assert rejected_text == original_text, (
        f"Rejected text mismatch.\n"
        f"  Expected (original): {original_text!r}\n"
        f"  Got (rejected):      {rejected_text!r}"
    )

    # Cleanup
    parsed_doc.package.close()


# ---------------------------------------------------------------------------
# Hypothesis strategies — Property 2 (formatting preservation)
# ---------------------------------------------------------------------------


@st.composite
def formatting_options(draw: st.DrawFn) -> dict[str, Any]:
    """Génère un ensemble de propriétés de formatage pour un run.

    Returns:
        Dictionnaire avec les options de formatage (au moins une active).
    """
    bold = draw(st.booleans())
    italic = draw(st.booleans())
    underline = draw(st.booleans())
    font_size = draw(st.sampled_from([None, 20, 24, 28, 32, 48]))
    font_color = draw(
        st.sampled_from([None, "FF0000", "0000FF", "00FF00", "333333"])
    )

    # S'assurer qu'au moins une propriété de formatage est active
    if not any([bold, italic, underline, font_size, font_color]):
        bold = True

    return {
        "bold": bold,
        "italic": italic,
        "underline": underline,
        "font_size": font_size,
        "font_color": font_color,
    }


@st.composite
def formatted_docx_with_correction(
    draw: st.DrawFn,
) -> tuple[bytes, str, str, dict[str, Any]]:
    """Génère un .docx avec formatage varié et une correction associée.

    Returns:
        Tuple (file_bytes, original_text, corrected_text, formatting)
        où formatting décrit les propriétés appliquées.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    fmt = draw(formatting_options())

    # Générer un texte original avec au moins 2 mots pour permettre un diff
    words = draw(
        st.lists(
            st.text(
                alphabet=st.characters(whitelist_categories=("L", "N")),
                min_size=2,
                max_size=15,
            ),
            min_size=2,
            max_size=8,
        )
    )
    original_text = " ".join(words)

    # Créer un paragraphe avec un run formaté
    para = doc.add_paragraph()
    run = para.add_run(original_text)

    # Appliquer le formatage
    run.bold = fmt["bold"]
    run.italic = fmt["italic"]
    if fmt["underline"]:
        run.underline = True
    if fmt["font_size"] is not None:
        run.font.size = Pt(fmt["font_size"])
    if fmt["font_color"] is not None:
        run.font.color.rgb = RGBColor.from_string(fmt["font_color"])

    # Générer un texte corrigé différent de l'original
    # Remplacer le premier mot par un mot différent
    replacement_word = draw(
        st.text(
            alphabet=st.characters(whitelist_categories=("L",)),
            min_size=2,
            max_size=15,
        ).filter(lambda w: w != words[0])
    )
    corrected_words = [replacement_word] + words[1:]
    corrected_text = " ".join(corrected_words)

    buffer = io.BytesIO()
    doc.save(buffer)
    file_bytes = buffer.getvalue()

    return file_bytes, original_text, corrected_text, fmt


# ---------------------------------------------------------------------------
# Helpers — Property 2 (formatting verification)
# ---------------------------------------------------------------------------


def _has_formatting_property(
    rpr_element: "etree._Element", fmt: dict[str, Any]
) -> dict[str, bool]:
    """Vérifie quelles propriétés de formatage sont présentes dans un w:rPr.

    Args:
        rpr_element: Élément XML w:rPr à inspecter.
        fmt: Dictionnaire des propriétés attendues.

    Returns:
        Dictionnaire indiquant quelles propriétés attendues sont présentes.
    """
    results: dict[str, bool] = {}

    if fmt["bold"]:
        b_elem = rpr_element.find(f"{{{WORD_NS}}}b")
        results["bold"] = b_elem is not None

    if fmt["italic"]:
        i_elem = rpr_element.find(f"{{{WORD_NS}}}i")
        results["italic"] = i_elem is not None

    if fmt["underline"]:
        u_elem = rpr_element.find(f"{{{WORD_NS}}}u")
        results["underline"] = u_elem is not None

    if fmt["font_size"] is not None:
        sz_elem = rpr_element.find(f"{{{WORD_NS}}}sz")
        results["font_size"] = sz_elem is not None

    if fmt["font_color"] is not None:
        color_elem = rpr_element.find(f"{{{WORD_NS}}}color")
        results["font_color"] = color_elem is not None

    return results


def _check_revision_elements_have_formatting(
    xml_bytes: bytes,
    fmt: dict[str, Any],
) -> tuple[bool, str]:
    """Vérifie que les éléments de révision préservent le formatage.

    Parse le document .docx de sortie et vérifie que les éléments w:del
    et w:ins contiennent les propriétés de formatage attendues.

    Args:
        xml_bytes: Bytes du fichier .docx de sortie.
        fmt: Propriétés de formatage attendues.

    Returns:
        Tuple (success, error_message).
    """
    target_para = _find_paragraph_with_revisions(xml_bytes)
    if target_para is None:
        return False, "No paragraph with w:del or w:ins elements found"

    # Chercher les éléments w:del et w:ins dans ce paragraphe
    del_elements = target_para.findall(f"{{{WORD_NS}}}del")
    ins_elements = target_para.findall(f"{{{WORD_NS}}}ins")

    if not del_elements and not ins_elements:
        return False, "No w:del or w:ins elements found in target paragraph"

    # Vérifier le formatage sur les runs dans w:del
    for del_elem in del_elements:
        runs = del_elem.findall(f".//{{{WORD_NS}}}r")
        for run_elem in runs:
            rpr = run_elem.find(f"{{{WORD_NS}}}rPr")
            if rpr is None:
                return False, (
                    "w:del run missing w:rPr element — "
                    "formatting not preserved on deletion"
                )
            check = _has_formatting_property(rpr, fmt)
            for prop_name, found in check.items():
                if not found:
                    return False, (
                        f"w:del run missing formatting property: {prop_name}"
                    )

    # Vérifier le formatage sur les runs dans w:ins
    for ins_elem in ins_elements:
        runs = ins_elem.findall(f".//{{{WORD_NS}}}r")
        for run_elem in runs:
            rpr = run_elem.find(f"{{{WORD_NS}}}rPr")
            if rpr is None:
                return False, (
                    "w:ins run missing w:rPr element — "
                    "formatting not preserved on insertion"
                )
            check = _has_formatting_property(rpr, fmt)
            for prop_name, found in check.items():
                if not found:
                    return False, (
                        f"w:ins run missing formatting property: {prop_name}"
                    )

    return True, ""


# ---------------------------------------------------------------------------
# Property 2 : Préservation du formatage des runs
# ---------------------------------------------------------------------------


@settings(max_examples=100, deadline=None)
@given(data=formatted_docx_with_correction())
def test_run_formatting_preservation(
    data: tuple[bytes, str, str, dict[str, Any]],
) -> None:
    """Pour tout paragraphe avec formatage, le TrackChangesGenerator préserve
    les propriétés de formatage sur les éléments w:del et w:ins.

    **Validates: Requirements 3.2, 5.4**
    """
    file_bytes, original_text, corrected_text, fmt = data

    # 1. Parser le document
    parser = DocumentParser()
    parsed_doc = parser.parse(file_bytes, "docx")

    # 2. Trouver le paragraphe contenant le texte original
    target_para = None
    for para in parsed_doc.paragraphs:
        if para.full_text == original_text:
            target_para = para
            break

    assert target_para is not None, (
        f"Could not find paragraph with text {original_text!r} in parsed doc"
    )

    # 3. Créer une correction pour le paragraphe cible
    correction = ParagraphCorrection(
        paragraph_index=target_para.index,
        original_text=original_text,
        corrected_text=corrected_text,
        has_changes=True,
    )

    # 4. Générer les track changes
    generator = TrackChangesGenerator()
    output_bytes = generator.generate(parsed_doc, [correction])

    # 5. Vérifier que le formatage est préservé sur w:del et w:ins
    success, error_msg = _check_revision_elements_have_formatting(
        output_bytes, fmt
    )

    # Cleanup
    parsed_doc.package.close()

    assert success, (
        f"Formatting preservation failed: {error_msg}\n"
        f"Original: {original_text!r}\n"
        f"Corrected: {corrected_text!r}\n"
        f"Formatting: {fmt}"
    )


# ---------------------------------------------------------------------------
# Property 3 : Validité structurelle des éléments de révision
# ---------------------------------------------------------------------------


@settings(max_examples=100, deadline=None)
@given(data=paragraph_correction_pairs())
def test_revision_element_validity(
    data: tuple[bytes, str, str],
) -> None:
    """Pour toute paire (original, corrigé) où les textes diffèrent, le XML
    généré DOIT contenir des éléments de révision OOXML valides : chaque w:del
    et w:ins possède les attributs w:id, w:author et w:date, tous les w:id
    sont uniques dans le document, et w:author vaut "Judi-Expert".

    **Validates: Requirements 5.2, 5.5**
    """
    docx_bytes, original_text, corrected_text = data

    # 1. Parser le document
    parser = DocumentParser()
    parsed_doc = parser.parse(docx_bytes, "docx")

    # 2. Trouver le paragraphe contenant le texte original
    target_para = None
    for para in parsed_doc.paragraphs:
        if para.full_text == original_text:
            target_para = para
            break

    assert target_para is not None, (
        f"Could not find paragraph with text {original_text!r} in parsed doc"
    )

    # 3. Créer la correction
    correction = ParagraphCorrection(
        paragraph_index=target_para.index,
        original_text=original_text,
        corrected_text=corrected_text,
        has_changes=True,
    )

    # 4. Générer les track changes
    generator = TrackChangesGenerator()
    output_bytes = generator.generate(parsed_doc, [correction])

    # 5. Extraire le XML du document de sortie
    with zipfile.ZipFile(io.BytesIO(output_bytes), "r") as zf:
        doc_xml = zf.read("word/document.xml")

    root = etree.fromstring(doc_xml)

    # 6. Trouver tous les éléments w:del et w:ins dans le document
    all_del = root.findall(f".//{{{WORD_NS}}}del")
    all_ins = root.findall(f".//{{{WORD_NS}}}ins")

    revision_elements = all_del + all_ins

    # Il doit y avoir au moins un élément de révision puisque les textes diffèrent
    assert len(revision_elements) > 0, (
        "No revision elements (w:del or w:ins) found despite texts being different.\n"
        f"  Original: {original_text!r}\n"
        f"  Corrected: {corrected_text!r}"
    )

    # 7. Vérifier les attributs requis et collecter les w:id
    all_ids: list[str] = []
    w_id_attr = f"{{{WORD_NS}}}id"
    w_author_attr = f"{{{WORD_NS}}}author"
    w_date_attr = f"{{{WORD_NS}}}date"

    for elem in revision_elements:
        tag_local = etree.QName(elem.tag).localname

        # Vérifier w:id
        rev_id = elem.get(w_id_attr)
        assert rev_id is not None, (
            f"Element w:{tag_local} is missing required attribute w:id"
        )
        all_ids.append(rev_id)

        # Vérifier w:author
        author = elem.get(w_author_attr)
        assert author is not None, (
            f"Element w:{tag_local} (id={rev_id}) is missing required attribute w:author"
        )
        assert author == "Judi-Expert", (
            f"Element w:{tag_local} (id={rev_id}) has w:author={author!r}, "
            f"expected 'Judi-Expert'"
        )

        # Vérifier w:date
        date = elem.get(w_date_attr)
        assert date is not None, (
            f"Element w:{tag_local} (id={rev_id}) is missing required attribute w:date"
        )

    # 8. Vérifier l'unicité des w:id
    id_set = set(all_ids)
    assert len(id_set) == len(all_ids), (
        f"Duplicate w:id values found. "
        f"Total IDs: {len(all_ids)}, Unique IDs: {len(id_set)}. "
        f"Duplicates: {[x for x in all_ids if all_ids.count(x) > 1]}"
    )

    # Cleanup
    parsed_doc.package.close()
