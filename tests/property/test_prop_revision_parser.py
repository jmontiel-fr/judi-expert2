"""Property-based tests — Document Revision Parser.

Feature: document-revision

Property-based tests validating correctness properties of the
DocumentParser component for the document revision feature.

- Property 1: Préservation du nombre de paragraphes (Validates: Requirements 3.1, 5.4)
- Property 6: Rejet des fichiers invalides (Validates: Requirements 2.2, 7.2, 7.5)
- Property 7: Idempotence du parsing sans correction (Validates: Requirements 3.1, 3.4)
"""

import io
import sys
import zipfile
from pathlib import Path

from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

# Ajouter le backend au path pour importer les services
sys.path.insert(
    0,
    str(
        Path(__file__).resolve().parents[2]
        / "local-site"
        / "web"
        / "backend"
    ),
)

from services.document_parser import DocumentParser
from services.revision_models import DocumentParseError
from services.track_changes_generator import TrackChangesGenerator


# ---------------------------------------------------------------------------
# Hypothesis strategies
# ---------------------------------------------------------------------------


@st.composite
def docx_documents(draw: st.DrawFn) -> bytes:
    """Génère des documents .docx valides avec contenu aléatoire.

    Produit des documents contenant entre 1 et 20 paragraphes, chacun
    composé de 1 à 5 runs avec formatage varié (bold, italic).
    """
    doc = Document()
    n_paragraphs = draw(st.integers(min_value=1, max_value=20))

    for _ in range(n_paragraphs):
        text = draw(
            st.text(
                alphabet=st.characters(
                    whitelist_categories=("L", "N", "P", "Z"),
                    blacklist_characters="\x00",
                ),
                min_size=1,
                max_size=200,
            )
        )
        para = doc.add_paragraph()
        # Ajouter des runs avec formatage varié
        n_runs = draw(st.integers(min_value=1, max_value=5))
        words = text.split() or [text]
        for i in range(n_runs):
            run_text = words[i % len(words)] if words else "x"
            run = para.add_run(run_text + " ")
            run.bold = draw(st.booleans())
            run.italic = draw(st.booleans())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Propriété 1 — Préservation du nombre de paragraphes
# ---------------------------------------------------------------------------


@settings(max_examples=100, deadline=None)
@given(doc_bytes=docx_documents())
def test_paragraph_count_preservation(doc_bytes: bytes) -> None:
    """Pour tout document .docx valide, le parsing préserve le nombre exact
    de paragraphes présents dans le body XML du document.

    **Validates: Requirements 3.1, 5.4**

    Vérifie que le nombre de ParagraphInfo retournés par le parser correspond
    exactement au nombre d'éléments w:p dans le body du document source.
    """
    from lxml import etree

    parser = DocumentParser()

    # Parser le document
    parsed_doc = parser.parse(doc_bytes, "docx")

    # Compter les paragraphes directement dans le XML source pour référence
    package = zipfile.ZipFile(io.BytesIO(doc_bytes), "r")
    document_xml = package.read("word/document.xml")
    tree = etree.fromstring(document_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body = tree.find(".//w:body", ns)
    xml_paragraph_count = len(body.findall("w:p", ns))
    package.close()

    assert len(parsed_doc.paragraphs) == xml_paragraph_count, (
        f"Expected {xml_paragraph_count} paragraphs (from XML), "
        f"got {len(parsed_doc.paragraphs)}"
    )

    # Cleanup: fermer le package ZIP
    parsed_doc.package.close()


# ---------------------------------------------------------------------------
# Propriété 7 — Idempotence du parsing sans correction
# ---------------------------------------------------------------------------


@settings(max_examples=100, deadline=None)
@given(doc_bytes=docx_documents())
def test_parsing_idempotence_without_corrections(doc_bytes: bytes) -> None:
    """Pour tout document .docx valide, parse → serialize (sans corrections)
    → re-parse doit produire un contenu textuel identique à l'original.

    Vérifie que le round-trip parse/serialize préserve fidèlement le
    contenu textuel de chaque paragraphe lorsqu'aucune correction n'est
    appliquée.
    """
    parser = DocumentParser()
    generator = TrackChangesGenerator()

    # 1. Parser le document original
    parsed_original = parser.parse(doc_bytes, "docx")
    original_texts = [p.full_text for p in parsed_original.paragraphs]

    # 2. Sérialiser sans corrections (liste vide)
    output_bytes = generator.generate(parsed_original, corrections=[])

    # 3. Re-parser le document sérialisé
    parsed_output = parser.parse(output_bytes, "docx")
    output_texts = [p.full_text for p in parsed_output.paragraphs]

    # 4. Vérifier que le contenu textuel est identique
    assert len(output_texts) == len(original_texts), (
        f"Nombre de paragraphes différent : "
        f"original={len(original_texts)}, output={len(output_texts)}"
    )

    for i, (orig, out) in enumerate(zip(original_texts, output_texts)):
        assert orig == out, (
            f"Paragraphe {i} différent après round-trip :\n"
            f"  original : {orig!r}\n"
            f"  output   : {out!r}"
        )


# ---------------------------------------------------------------------------
# Strategies for Property 6
# ---------------------------------------------------------------------------


@st.composite
def non_zip_bytes(draw: st.DrawFn) -> bytes:
    """Generate arbitrary byte sequences that are NOT valid ZIP archives.

    Filters out any bytes that happen to start with the ZIP magic number
    (PK\\x03\\x04) to ensure they cannot be valid ZIP files.
    """
    data = draw(st.binary(min_size=0, max_size=1024))
    # Ensure the bytes don't accidentally form a valid ZIP
    if data[:4] == b"PK\x03\x04":
        # Corrupt the magic number
        data = b"\x00" + data[1:]
    return data


@st.composite
def zip_without_document_xml(draw: st.DrawFn) -> bytes:
    """Generate valid ZIP archives that do NOT contain word/document.xml.

    Creates ZIP files with arbitrary entries but ensures none of them
    is named 'word/document.xml'.
    """
    num_entries = draw(st.integers(min_value=0, max_value=5))
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in range(num_entries):
            # Generate filenames that are NOT word/document.xml
            name = draw(
                st.text(
                    alphabet=st.characters(
                        whitelist_categories=("L", "N"),
                        whitelist_characters="_-/.",
                    ),
                    min_size=1,
                    max_size=30,
                ).filter(lambda n: n != "word/document.xml")
            )
            content = draw(st.binary(min_size=0, max_size=256))
            zf.writestr(name, content)
    return buffer.getvalue()


@st.composite
def non_utf8_bytes(draw: st.DrawFn) -> bytes:
    """Generate byte sequences that are NOT valid UTF-8.

    Produces bytes containing invalid UTF-8 sequences to test
    the parser's rejection of non-UTF-8 encoded text files.
    """
    # Start with some valid bytes, then inject invalid UTF-8 sequences
    prefix = draw(st.binary(min_size=0, max_size=128))
    # Invalid UTF-8 sequences: continuation byte without start, overlong, etc.
    invalid_sequences = [
        b"\x80",           # Lone continuation byte
        b"\xc0\xaf",      # Overlong encoding
        b"\xfe\xff",      # Invalid start bytes
        b"\xff\xfe",      # Invalid start bytes
        b"\xc0",          # Truncated 2-byte sequence
        b"\xe0\x80",      # Truncated 3-byte sequence
        b"\xf0\x80\x80",  # Truncated 4-byte sequence
        b"\xed\xa0\x80",  # Surrogate half (U+D800)
        b"\xed\xbf\xbf",  # Surrogate half (U+DFFF)
    ]
    invalid = draw(st.sampled_from(invalid_sequences))
    suffix = draw(st.binary(min_size=0, max_size=128))
    result = prefix + invalid + suffix
    # Verify it's actually not valid UTF-8
    try:
        result.decode("utf-8")
        # If it decoded successfully, force an invalid byte
        result = result + b"\xff"
    except UnicodeDecodeError:
        pass
    return result


# ---------------------------------------------------------------------------
# Propriété 6 — Rejet des fichiers invalides
# ---------------------------------------------------------------------------


class TestProperty6InvalidFileRejection:
    """Property 6: Rejet des fichiers invalides.

    For any byte sequence that is not a valid .docx file (not a ZIP archive,
    or ZIP without word/document.xml), the DocumentParser SHALL raise a
    DocumentParseError. For any byte sequence that is not valid UTF-8,
    parsing as .txt or .md SHALL raise DocumentParseError.

    **Validates: Requirements 2.2, 7.2, 7.5**
    """

    @given(data=non_zip_bytes())
    @settings(max_examples=100, deadline=None)
    def test_non_zip_bytes_raise_parse_error(self, data: bytes) -> None:
        """Non-ZIP byte sequences must raise DocumentParseError for .docx.

        **Validates: Requirements 2.2, 7.2**
        """
        parser = DocumentParser()
        try:
            parser.parse(data, "docx")
            raise AssertionError(
                "DocumentParser.parse() did not raise DocumentParseError "
                "for non-ZIP bytes"
            )
        except DocumentParseError:
            pass  # Expected behavior

    @given(data=zip_without_document_xml())
    @settings(max_examples=100, deadline=None)
    def test_zip_without_document_xml_raises_parse_error(
        self, data: bytes
    ) -> None:
        """ZIP archives without word/document.xml must raise DocumentParseError.

        **Validates: Requirements 2.2, 7.2**
        """
        parser = DocumentParser()
        try:
            parser.parse(data, "docx")
            raise AssertionError(
                "DocumentParser.parse() did not raise DocumentParseError "
                "for ZIP without word/document.xml"
            )
        except DocumentParseError:
            pass  # Expected behavior

    @given(data=non_utf8_bytes())
    @settings(max_examples=100, deadline=None)
    def test_non_utf8_bytes_txt_raises_parse_error(
        self, data: bytes
    ) -> None:
        """Non-UTF-8 byte sequences must raise DocumentParseError for .txt.

        **Validates: Requirements 7.5**
        """
        parser = DocumentParser()
        try:
            parser.parse(data, "txt")
            raise AssertionError(
                "DocumentParser.parse() did not raise DocumentParseError "
                "for non-UTF-8 bytes as .txt"
            )
        except DocumentParseError:
            pass  # Expected behavior

    @given(data=non_utf8_bytes())
    @settings(max_examples=100, deadline=None)
    def test_non_utf8_bytes_md_raises_parse_error(
        self, data: bytes
    ) -> None:
        """Non-UTF-8 byte sequences must raise DocumentParseError for .md.

        **Validates: Requirements 7.5**
        """
        parser = DocumentParser()
        try:
            parser.parse(data, "md")
            raise AssertionError(
                "DocumentParser.parse() did not raise DocumentParseError "
                "for non-UTF-8 bytes as .md"
            )
        except DocumentParseError:
            pass  # Expected behavior
