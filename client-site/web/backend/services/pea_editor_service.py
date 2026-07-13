"""
Judi-Expert — Service de parsing PEA/TPE.

Parse un document .docx PEA ou TPE pour extraire sa structure en blocs
éditables : headings, texte, placeholders et annotations.

Gère les annotations mono-ligne et multi-paragraphes, la détection de
headings par style Word et par numérotation, et la classification
d'éditabilité des blocs.

Valide : Exigences 6.2, 6.3, 8.2, 8.3, 11.1, 11.2, 11.3, 11.4, 11.6, 11.7
"""

from __future__ import annotations

import io
import logging
import re

from docx import Document as DocxDocument

from .pea_editor_models import (
    AnnotationBlock,
    HeadingBlock,
    PEABlock,
    PEADocument,
    PlaceholderBlock,
    SectionInfo,
    TextBlock,
)

logger = logging.getLogger(__name__)


class PEAEditorService:
    """Parse un document PEA/TPE en blocs éditables.

    Extrait les headings, placeholders, annotations et texte normal
    d'un fichier .docx pour les présenter sous forme de formulaire.
    """

    ANNOTATION_TYPES: frozenset = frozenset(
        {
            "remplir",
            "dires",
            "analyse",
            "conclusion",
            "verbatim",
            "resume",
            "reference",
            "cite",
            "question",
            "age",
        }
    )

    # Types éditables dans le formulaire PEA
    EDITABLE_TYPES: frozenset = frozenset(
        {"remplir", "dires", "analyse", "conclusion"}
    )

    # Types traités comme @remplir_champ (champs inline éditables)
    FIELD_TYPES: frozenset = frozenset(
        {"remplir"}
    )

    # Regex patterns
    # Heading numbering: x.x.x (1-4 levels)
    HEADING_NUMBER_RE = re.compile(
        r"^(\d+(?:\.\d+){0,3})\s+(.+)$"
    )

    # Placeholder: <<name>>
    PLACEHOLDER_RE = re.compile(r"<<([^>]+)>>")

    # Annotation inline: @type_suffix contenu@ or @type@ anywhere in text
    # Le type doit commencer par une lettre et le @ doit être précédé d'un espace ou en début
    # Matches both @type contenu@ and @type@ (no content)
    ANNOTATION_INLINE_RE = re.compile(
        r"(?:^|(?<=\s))@(\w+(?:_[\w.]+)*)\s*(.*?)\s*@(?=\s|$|[.,;:!?)])"
    )

    # Annotation single-line: entire line is one annotation
    ANNOTATION_SINGLE_RE = re.compile(
        r"^\s*@(\w+(?:_[\w.]+)*)\s*(.*?)\s*@\s*$", re.DOTALL
    )

    # Annotation opening: @type_suffix contenu... (no closing @ on same line)
    ANNOTATION_OPEN_RE = re.compile(r"^\s*@(\w+(?:_[\w.]+)*)\s+(.*)")

    # Annotation close: line ending with @ (standalone or at end)
    ANNOTATION_CLOSE_RE = re.compile(r"^(.*?)\s*@\s*$", re.DOTALL)

    # Section reference for @dires/@analyse: section_x.x.x
    SECTION_REF_RE = re.compile(
        r"^section_(\d+(?:\.\d+){0,3})\s*(.*)$"
    )

    # Field name for @remplir: field_name format :
    REMPLIR_RE = re.compile(
        r"^([a-z][a-z0-9_]{0,63})\s*(.*?)\s*:\s*(.*)$", re.DOTALL
    )

    # Word heading styles
    HEADING_STYLES = frozenset(
        {
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Heading 4",
            "Heading 5",
            "Heading 6",
            "Heading1",
            "Heading2",
            "Heading3",
            "Heading4",
            "Heading5",
            "Heading6",
            "Titre 1",
            "Titre 2",
            "Titre 3",
            "Titre 4",
            "Titre 5",
            "Titre 6",
            "Titre1",
            "Titre2",
            "Titre3",
            "Titre4",
            "Titre5",
            "Titre6",
            "Title",
            "Titre",
        }
    )

    def parse(self, file_bytes: bytes) -> PEADocument:
        """Parse le .docx et retourne un PEADocument structuré.

        Args:
            file_bytes: Contenu binaire du fichier .docx.

        Returns:
            PEADocument contenant les blocs, sections et erreurs.

        Raises:
            ValueError: Si le fichier n'est pas un .docx valide.
        """
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            raise ValueError(
                f"Le fichier .docx est invalide ou corrompu : {e}"
            ) from e

        blocks, parse_errors = self._extract_blocks(doc)
        sections = self._build_sections(blocks)
        errors = parse_errors

        return PEADocument(
            blocks=blocks,
            sections=sections,
            errors=errors,
            total_paragraphs=len(doc.paragraphs),
        )

    def _extract_blocks(
        self, doc: DocxDocument
    ) -> tuple[list[PEABlock], list[str]]:
        """Extrait les blocs ordonnés du document.

        Parcourt chaque paragraphe et détecte :
        - Headings (par style Word ou numérotation x.x.x)
        - Placeholders <<...>>
        - Annotations @type...@
        - Texte normal

        Gère les annotations multi-paragraphes.

        Args:
            doc: Document python-docx ouvert.

        Returns:
            Tuple (liste ordonnée de PEABlock, liste d'erreurs de parsing).
        """
        blocks: list[PEABlock] = []
        errors: list[str] = []
        open_annotation: dict | None = None

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text

            # If a multi-paragraph annotation is open
            if open_annotation is not None:
                close_match = self.ANNOTATION_CLOSE_RE.match(text)
                if close_match:
                    # Close the annotation
                    remaining = close_match.group(1)
                    if remaining:
                        open_annotation["content_parts"].append(remaining)
                    full_content = "\n".join(open_annotation["content_parts"])
                    annotation = self._build_annotation_block(
                        annotation_type=open_annotation["type"],
                        suffix_raw=open_annotation["suffix_raw"],
                        content=full_content.strip(),
                        paragraph_index=open_annotation["position"],
                    )
                    blocks.append(annotation)
                    open_annotation = None
                else:
                    # Continue accumulating content
                    open_annotation["content_parts"].append(text)
                continue

            # 1. Detect heading
            heading = self._detect_heading(paragraph, idx)
            if heading is not None:
                blocks.append(heading)
                continue

            # 2. Check if line contains annotations (inline or full-line)
            # Only attempt annotation detection if line contains @
            if "@" not in text:
                # No @ at all — check for placeholders or plain text
                if self.PLACEHOLDER_RE.search(text):
                    match = self.PLACEHOLDER_RE.search(text)
                    blocks.append(
                        PlaceholderBlock(
                            paragraph_index=idx,
                            name=match.group(1),
                            full_text=text,
                        )
                    )
                elif text.strip():
                    blocks.append(TextBlock(paragraph_index=idx, content=text))
                else:
                    blocks.append(TextBlock(paragraph_index=idx, content=""))
                continue

            inline_annotations = list(self.ANNOTATION_INLINE_RE.finditer(text))

            # Filtrer les faux positifs : ne garder que les matches dont le type
            # est connu ou ressemble à un identifiant d'annotation valide
            valid_annotations = []
            for match in inline_annotations:
                full_key = match.group(1).lower()
                base_type, _ = self._split_annotation_key(full_key)
                # Accepter si le type de base est connu OU si le mot complet
                # ressemble à un identifiant (contient un underscore ou est court)
                if (base_type in self.ANNOTATION_TYPES or
                    "_" in full_key or
                    len(full_key) <= 15):
                    valid_annotations.append(match)
            inline_annotations = valid_annotations

            if inline_annotations:
                # Line has one or more complete annotations (@...@)
                # Extract text before/between/after annotations as text/placeholder blocks
                last_end = 0
                for match in inline_annotations:
                    # Text before this annotation
                    before = text[last_end:match.start()].strip()
                    if before:
                        self._emit_residual_text(before, idx, blocks)

                    # The annotation itself
                    full_key = match.group(1).lower()
                    rest = match.group(2)
                    ann_type, key_suffix = self._split_annotation_key(full_key)
                    suffix_raw = f"{key_suffix} {rest}".strip() if key_suffix else rest
                    annotation = self._build_annotation_block(
                        annotation_type=ann_type,
                        suffix_raw=suffix_raw,
                        content="",
                        paragraph_index=idx,
                    )
                    blocks.append(annotation)
                    last_end = match.end()

                # Text after last annotation
                after = text[last_end:].strip()
                if after:
                    self._emit_residual_text(after, idx, blocks)
                continue

            # 3. Check for annotation opening (multi-paragraph, no closing @ on line)
            # Only if no complete annotation was found above
            # A line is a multi-paragraph opening if it starts with @type but has no closing @
            if text.strip().startswith("@") and "@" not in text.strip()[1:].rstrip():
                open_match = self.ANNOTATION_OPEN_RE.match(text)
                if open_match:
                    full_key = open_match.group(1).lower()
                    rest = open_match.group(2)
                    ann_type, key_suffix = self._split_annotation_key(full_key)
                    suffix_raw = f"{key_suffix} {rest}".strip() if key_suffix else rest
                    open_annotation = {
                        "type": ann_type,
                        "suffix_raw": suffix_raw,
                        "position": idx,
                        "content_parts": [],
                    }
                    continue

            # 4. Detect placeholder <<...>> (no annotation on this line)
            if self.PLACEHOLDER_RE.search(text):
                match = self.PLACEHOLDER_RE.search(text)
                blocks.append(
                    PlaceholderBlock(
                        paragraph_index=idx,
                        name=match.group(1),
                        full_text=text,
                    )
                )
                continue

            # 5. Normal text
            if text.strip():
                blocks.append(
                    TextBlock(
                        paragraph_index=idx,
                        content=text,
                    )
                )
            else:
                # Empty paragraphs are still text blocks
                blocks.append(
                    TextBlock(
                        paragraph_index=idx,
                        content="",
                    )
                )

        # Handle unclosed annotation at end of document
        if open_annotation is not None:
            errors.append(
                f"Paragraphe {open_annotation['position']} : annotation "
                f"@{open_annotation['type']} non fermée (@ manquant)."
            )
            # Still include the partial annotation in blocks
            full_content = "\n".join(open_annotation["content_parts"])
            annotation = self._build_annotation_block(
                annotation_type=open_annotation["type"],
                suffix_raw=open_annotation["suffix_raw"],
                content=full_content.strip(),
                paragraph_index=open_annotation["position"],
            )
            blocks.append(annotation)

        return blocks, errors

    def _detect_heading(self, paragraph, paragraph_index: int) -> HeadingBlock | None:
        """Détecte si un paragraphe est un titre de section.

        Détection par :
        1. Style Word (Heading 1-6, Titre 1-6, etc.)
        2. Regex de numérotation x.x.x au début du texte

        Args:
            paragraph: Paragraphe python-docx.
            paragraph_index: Index du paragraphe dans le document.

        Returns:
            HeadingBlock si c'est un heading, None sinon.
        """
        text = paragraph.text.strip()
        if not text:
            return None

        # Check by Word style
        style_name = ""
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name

        if style_name in self.HEADING_STYLES:
            level = self._style_to_level(style_name)
            number, title = self._extract_heading_parts(text)
            # Si pas de numéro dans le texte, essayer d'extraire depuis la numérotation Word
            if not number:
                number = self._extract_word_numbering(paragraph)
            return HeadingBlock(
                paragraph_index=paragraph_index,
                level=level,
                number=number,
                text=title,
            )

        # Check by numbering regex (e.g., "2.1.3 Titre de section")
        num_match = self.HEADING_NUMBER_RE.match(text)
        if num_match:
            number = num_match.group(1)
            title = num_match.group(2)
            # Determine level from number of dots + 1
            level = number.count(".") + 1
            return HeadingBlock(
                paragraph_index=paragraph_index,
                level=level,
                number=number,
                text=title,
            )

        return None

    def _style_to_level(self, style_name: str) -> int:
        """Convertit un nom de style Word en niveau de heading (1-6).

        Args:
            style_name: Nom du style Word.

        Returns:
            Niveau de heading (1-6).
        """
        # Extract digit from style name
        for char in style_name:
            if char.isdigit():
                level = int(char)
                if 1 <= level <= 6:
                    return level
        # "Title" or "Titre" without number → level 1
        return 1

    def _split_annotation_key(self, full_key: str) -> tuple[str, str]:
        """Sépare la clé d'annotation en type de base et suffix.

        Le format est @type_suffix où type est un mot-clé connu
        (remplir, dires, analyse, etc.) et suffix est le reste.

        Exemples:
            "remplir_date_entretien" → ("remplir", "date_entretien")
            "dires_3.1.2" → ("dires", "3.1.2")
            "analyse_2.1" → ("analyse", "2.1")
            "conclusion" → ("conclusion", "")
            "verbatim" → ("verbatim", "")
            "resume_3.2" → ("resume", "3.2")

        Args:
            full_key: Clé complète capturée par le regex.

        Returns:
            Tuple (type_de_base, suffix).
        """
        # Essayer de matcher un type connu au début
        for known_type in sorted(self.ANNOTATION_TYPES, key=len, reverse=True):
            if full_key == known_type:
                return known_type, ""
            if full_key.startswith(known_type + "_"):
                suffix = full_key[len(known_type) + 1:]
                return known_type, suffix

        # Type inconnu : garder le type complet (pas de split)
        # Ex: "date_naissance_pex" → ("date_naissance_pex", "")
        return full_key, ""

    def _extract_heading_parts(self, text: str) -> tuple[str, str]:
        """Extrait le numéro et le titre d'un heading.

        Si le texte commence par une numérotation x.x.x, on la sépare.
        Sinon, le numéro est vide et le titre est le texte complet.

        Args:
            text: Texte du paragraphe.

        Returns:
            Tuple (number, title).
        """
        num_match = self.HEADING_NUMBER_RE.match(text)
        if num_match:
            return num_match.group(1), num_match.group(2)
        return "", text

    def _extract_word_numbering(self, paragraph) -> str:
        """Extrait la numérotation automatique Word d'un paragraphe.

        Cherche dans les propriétés XML du paragraphe (numPr) et dans
        le texte des runs pour trouver un numéro de section.

        Args:
            paragraph: Paragraphe python-docx.

        Returns:
            Numéro de section trouvé, ou chaîne vide.
        """
        try:
            # Méthode 1 : Chercher dans les propriétés de numérotation XML
            pPr = paragraph._element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if pPr is not None:
                # La numérotation existe mais le texte n'est pas dans paragraph.text
                # Essayer de reconstruire depuis les runs
                pass

            # Méthode 2 : Chercher un numéro dans le premier run
            # (certains documents mettent le numéro dans un run séparé)
            if paragraph.runs:
                first_run_text = paragraph.runs[0].text.strip()
                num_match = re.match(r"^(\d+(?:\.\d+){0,3})\s*$", first_run_text)
                if num_match:
                    return num_match.group(1)

        except Exception:
            pass

        return ""

    # Regex to detect annotation opening at start of residual text
    _ANNOTATION_START_RE = re.compile(r"^@(\w+(?:_[\w.]+)*)\s*(.*)")

    def _emit_residual_text(
        self,
        text: str,
        paragraph_index: int,
        blocks: list[PEABlock],
    ) -> None:
        """Handle residual text between/after inline annotations.

        If the text starts with @type (annotation opening), treat it as
        an inline annotation (if it also closes with @) or as a text block.
        Only fall back to placeholder detection if no annotation pattern.

        Args:
            text: Residual text to process.
            paragraph_index: Index of the source paragraph.
            blocks: List to append new blocks to.
        """
        if not text:
            return

        # Check if residual text starts with an annotation pattern (@word...)
        ann_start = self._ANNOTATION_START_RE.match(text)
        if ann_start:
            full_key = ann_start.group(1).lower()
            rest = ann_start.group(2)
            base_type, key_suffix = self._split_annotation_key(full_key)

            # Check if it's a complete inline annotation (has closing @)
            # Try to match the whole residual as a single-line annotation
            single_match = self.ANNOTATION_SINGLE_RE.match(text)
            if single_match:
                s_key = single_match.group(1).lower()
                s_rest = single_match.group(2)
                s_type, s_suffix = self._split_annotation_key(s_key)
                suffix_raw = (
                    f"{s_suffix} {s_rest}".strip() if s_suffix else s_rest
                )
                annotation = self._build_annotation_block(
                    annotation_type=s_type,
                    suffix_raw=suffix_raw,
                    content="",
                    paragraph_index=paragraph_index,
                )
                blocks.append(annotation)
                return

            # It's an annotation opening without closing @ on this text
            # Treat as a multi-paragraph-style annotation captured in one block
            # (Word may merge paragraphs, so the closing @ is at the end)
            # Look for a closing @ anywhere in the rest
            if "@" in rest:
                # Find the last @ that could be a closer
                close_idx = rest.rfind("@")
                content_part = rest[:close_idx].strip()
                suffix_raw = (
                    f"{key_suffix} {content_part}".strip()
                    if key_suffix
                    else content_part
                )
                annotation = self._build_annotation_block(
                    annotation_type=base_type,
                    suffix_raw=suffix_raw,
                    content="",
                    paragraph_index=paragraph_index,
                )
                blocks.append(annotation)
                # Text after the closing @
                after_close = rest[close_idx + 1:].strip()
                if after_close:
                    self._emit_residual_text(
                        after_close, paragraph_index, blocks
                    )
                return

            # No closing @ — treat as annotation with all remaining as content
            suffix_raw = (
                f"{key_suffix} {rest}".strip() if key_suffix else rest
            )
            annotation = self._build_annotation_block(
                annotation_type=base_type,
                suffix_raw=suffix_raw,
                content="",
                paragraph_index=paragraph_index,
            )
            blocks.append(annotation)
            return

        # Not an annotation — check for placeholders
        ph_match = self.PLACEHOLDER_RE.search(text)
        if ph_match:
            blocks.append(
                PlaceholderBlock(
                    paragraph_index=paragraph_index,
                    name=ph_match.group(1),
                    full_text=text,
                )
            )
        else:
            blocks.append(
                TextBlock(paragraph_index=paragraph_index, content=text)
            )

    def _build_annotation_block(
        self,
        annotation_type: str,
        suffix_raw: str,
        content: str,
        paragraph_index: int,
    ) -> AnnotationBlock:
        """Construit un AnnotationBlock à partir des éléments parsés.

        Gère l'extraction spécifique selon le type :
        - @remplir : field_name, field_format
        - @dires/@analyse : section_ref
        - @conclusion : contenu multi-ligne

        Args:
            annotation_type: Type d'annotation (lowercase).
            suffix_raw: Texte brut après le type (avant le contenu).
            content: Contenu textuel de l'annotation.
            paragraph_index: Index du paragraphe.

        Returns:
            AnnotationBlock configuré.
        """
        is_editable = annotation_type in self.EDITABLE_TYPES
        suffix = suffix_raw.strip()
        field_name: str | None = None
        field_format: str | None = None
        section_ref: str | None = None
        final_content = content

        # Types inconnus (pas dans ANNOTATION_TYPES) → traiter comme champ éditable
        # Garde le type original pour l'affichage (ex: @date_naissance_pex)
        if annotation_type not in self.ANNOTATION_TYPES:
            is_editable = True
            field_name = annotation_type
            field_format = "champ"
            # Le suffix_raw contient le contenu pré-rempli
            if suffix_raw.strip():
                final_content = suffix_raw.strip()
            suffix = ""
            return AnnotationBlock(
                paragraph_index=paragraph_index,
                annotation_type=annotation_type,  # Garder le type original
                suffix=suffix,
                content=final_content,
                is_editable=is_editable,
                field_name=field_name,
                field_format=field_format,
                section_ref=section_ref,
            )

        if annotation_type == "remplir":
            # Nouveau format :
            # @remplir_champ texte@ → inline, field_name="champ", content="texte"
            # @remplir_bloc texte@ → bloc, field_name="bloc", content="texte"
            # Le suffix_raw contient "field_name [texte]" après _split_annotation_key
            stripped = suffix_raw.strip()

            # Extraire le field_name (premier mot)
            parts = stripped.split(None, 1)
            if parts:
                field_name = parts[0]
                inline_content = parts[1].strip() if len(parts) > 1 else ""
                # Si le contenu commence par ":", le retirer
                if inline_content.startswith(":"):
                    inline_content = inline_content[1:].strip()
                # Déterminer le format selon le field_name
                if field_name == "bloc":
                    field_format = "bloc"
                else:
                    field_format = "champ"
                if inline_content:
                    final_content = inline_content
                suffix = field_name
            else:
                # Pas de field_name → champ libre
                field_name = "libre"
                field_format = "champ"
                suffix = "libre"

        elif annotation_type in ("dires", "analyse"):
            # Format: suffix_raw = "3 .1.1_pere contenu..." ou "3.1.1_pere contenu..."
            # Le numéro peut être splitté : key_suffix="3" et rest=".1.1_pere contenu"
            # Il faut reconstituer le numéro complet
            stripped = suffix_raw.strip()

            # Reconstituer le numéro : chercher digits.digits[_suffix] au début
            # Le numéro peut commencer directement ou après recomposition
            # Pattern: optionnel point+chiffres répétés, optionnel _suffixe
            ref_match = re.match(
                r"^(\d+(?:\s*\.\s*\d+)*(?:\s*_\s*\w+)?)\s*(.*?)$",
                stripped,
                re.DOTALL,
            )
            if ref_match:
                raw_ref = ref_match.group(1)
                # Nettoyer les espaces dans le numéro de section
                section_ref = re.sub(r"\s+", "", raw_ref)
                rest = ref_match.group(2).strip()
                if rest.startswith(":"):
                    inline_content = rest[1:].strip()
                    if inline_content:
                        final_content = inline_content
                elif rest:
                    final_content = rest
                suffix = f"section_{section_ref}"
            elif stripped.startswith("section_"):
                # Format "section_x.x.x_suffixe : contenu"
                section_match = re.match(
                    r"^section_([\w.]+)\s*(.*?)$", stripped, re.DOTALL
                )
                if section_match:
                    section_ref = section_match.group(1)
                    rest = section_match.group(2).strip()
                    if rest.startswith(":"):
                        inline_content = rest[1:].strip()
                        if inline_content:
                            final_content = inline_content
                    suffix = f"section_{section_ref}"
                else:
                    suffix = stripped
            else:
                # Fallback : garder tel quel
                suffix = stripped

        elif annotation_type == "conclusion":
            # Parse: optional colon then content
            stripped = suffix_raw.strip()
            if stripped.startswith(":"):
                inline_content = stripped[1:].strip()
                if inline_content:
                    final_content = inline_content
                suffix = ""
            elif stripped:
                # Could be ":" alone or content
                suffix = stripped.rstrip(":")
            else:
                suffix = ""

        else:
            # Other types: extract content after optional colon
            stripped = suffix_raw.strip()
            if ":" in stripped:
                parts = stripped.split(":", 1)
                suffix = parts[0].strip()
                inline_content = parts[1].strip()
                if inline_content:
                    final_content = inline_content
            else:
                suffix = stripped

        return AnnotationBlock(
            paragraph_index=paragraph_index,
            annotation_type=annotation_type,
            suffix=suffix,
            content=final_content,
            is_editable=is_editable,
            field_name=field_name,
            field_format=field_format,
            section_ref=section_ref,
        )

    def _build_sections(self, blocks: list[PEABlock]) -> list[SectionInfo]:
        """Construit la liste des sections pour la palette d'annotations.

        Parcourt les blocs pour trouver les headings et les annotations
        @dires/@analyse associées.

        Args:
            blocks: Liste des blocs extraits.

        Returns:
            Liste de SectionInfo pour la palette.
        """
        sections: list[SectionInfo] = []

        for block in blocks:
            if isinstance(block, HeadingBlock) and block.number:
                # Default annotation type is "dires"
                sections.append(
                    SectionInfo(
                        number=block.number,
                        title=block.text,
                        level=block.level,
                        annotation_type="dires",
                    )
                )

        # Update annotation_type based on actual annotations found
        # Map section numbers to their SectionInfo
        section_map: dict[str, SectionInfo] = {
            s.number: s for s in sections
        }

        for block in blocks:
            if isinstance(block, AnnotationBlock) and block.section_ref:
                if block.section_ref in section_map:
                    section_map[block.section_ref].annotation_type = (
                        block.annotation_type
                    )

        return sections
