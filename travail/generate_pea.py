"""Génère le PEA (Plan d'Entretien Annoté) à partir du PE et du Pré-Rapport."""

from docx import Document
from docx.shared import Pt
from copy import deepcopy

# Lire le PE
pe = Document("travail/pe-test.docx")

# Lire le Pré-Rapport et extraire le contenu par section
rapport = Document("travail/Pre-Rapport-Test-original.docx")
rapport_text = "\n".join(p.text for p in rapport.paragraphs)

# Mapping des contenus du rapport vers les sections du PE
# On extrait les dires et analyses du rapport pour chaque section

sections_data = {}

# Parser le rapport pour extraire dires/analyses par section
current_section = ""
current_type = ""  # "dires" ou "analyse"
buffer = []

for p in rapport.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    
    # Détecter les sections
    text_lower = text.lower()
    
    if text_lower in ("relations familiales", "relations à ses parents"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "relations_familiales"
        buffer = []
        continue
    elif text_lower == "relations à sa fratrie":
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "fratrie"
        buffer = []
        continue
    elif text_lower.startswith("centres d'intérêt"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "loisirs"
        buffer = []
        continue
    elif text_lower.startswith("vie sociale"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "social"
        buffer = []
        continue
    elif text_lower.startswith("vie amoureuse"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "amour"
        buffer = []
        continue
    elif text_lower == "scolarité":
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "scolarite"
        buffer = []
        continue
    elif text_lower == "vie professionnelle":
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "pro"
        buffer = []
        continue
    elif text_lower == "avant les faits":
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "avant_faits"
        buffer = []
        continue
    elif text_lower == "pendant les faits":
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "pendant_faits"
        buffer = []
        continue
    elif text_lower.startswith("contexte de la révélation"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "revelation"
        buffer = []
        continue
    elif text_lower.startswith("après le dépôt"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "apres_plainte"
        buffer = []
        continue
    elif text_lower.startswith("étude des pièces"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "etude_pieces"
        buffer = []
        continue
    elif text_lower.startswith("relation au corps"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "sante"
        buffer = []
        continue
    elif text_lower.startswith("éléments cognitifs"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "cognitifs"
        buffer = []
        continue
    elif text_lower.startswith("éléments de personnalité"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "personnalite"
        buffer = []
        continue
    elif text_lower.startswith("analyse de la personne pendant"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "entretien"
        buffer = []
        continue
    elif text_lower.startswith("rapport à la sexualité"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "sexualite"
        buffer = []
        continue
    elif text_lower.startswith("symptômes d'intrusion"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "intrusion"
        buffer = []
        continue
    elif text_lower.startswith("symptômes d'évitement"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "evitement"
        buffer = []
        continue
    elif text_lower.startswith("symptômes d'altération"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "cognitions"
        buffer = []
        continue
    elif text_lower.startswith("symptômes d'éveil"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "eveil"
        buffer = []
        continue
    elif text_lower.startswith("symptômes dissociatifs"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_section = "dissociatifs"
        buffer = []
        continue
    
    # Détecter dires/analyse
    if text_lower.startswith("dires") or text_lower.startswith("dires :") or text_lower.startswith("dires:"):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_type = "dires"
        buffer = []
        # Si le texte contient plus que juste "Dires :"
        rest = text.split(":", 1)
        if len(rest) > 1 and rest[1].strip():
            buffer.append(rest[1].strip())
        continue
    elif text_lower.startswith("analyse") and (":" in text_lower[:10]):
        if buffer and current_section:
            sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)
        current_type = "analyse"
        buffer = []
        rest = text.split(":", 1)
        if len(rest) > 1 and rest[1].strip():
            buffer.append(rest[1].strip())
        continue
    
    # Accumuler le contenu
    if current_section and current_type:
        buffer.append(text)

# Flush final
if buffer and current_section:
    sections_data.setdefault(current_section, {}).setdefault(current_type, []).extend(buffer)

# Maintenant, créer le PEA en copiant le PE et en remplissant les annotations
pea = Document("travail/pe-test.docx")

# Mapping section PE → section rapport
pe_to_rapport = {
    "Relations familiales": "relations_familiales",
    "Relations à sa fratrie": "fratrie",
    "Centres d'intérêt et loisirs": "loisirs",
    "Vie sociale et amicale": "social",
    "Vie amoureuse et affective": "amour",
    "Scolarité": "scolarite",
    "Vie professionnelle": "pro",
    "Avant les faits": "avant_faits",
    "Pendant les faits": "pendant_faits",
    "Contexte de la révélation des faits": "revelation",
    "Après le dépôt de plainte": "apres_plainte",
    "ÉTUDE DES PIÈCES": "etude_pieces",
    "Relation au corps et à la santé": "sante",
    "Éléments cognitifs": "cognitifs",
    "Éléments de personnalité": "personnalite",
    "Analyse de la personne pendant l'entretien": "entretien",
    "Rapport à la sexualité": "sexualite",
    "Symptômes d'intrusion": "intrusion",
    "Symptômes d'évitement": "evitement",
    "Symptômes d'altération des cognitions et de l'humeur": "cognitions",
    "Symptômes d'éveil et de réactivité": "eveil",
    "Symptômes dissociatifs": "dissociatifs",
}

# Parcourir les paragraphes du PEA et injecter le contenu
current_pe_section = ""
i = 0
while i < len(pea.paragraphs):
    p = pea.paragraphs[i]
    text = p.text.strip()
    
    # Détecter la section courante
    for pe_name, rapport_key in pe_to_rapport.items():
        if text == pe_name:
            current_pe_section = rapport_key
            break
    
    # Remplacer @Dires par le contenu
    if text == "@Dires" and current_pe_section in sections_data:
        dires = sections_data[current_pe_section].get("dires", [])
        if dires:
            content = " ".join(dires)
            p.text = "@dires"
            # Ajouter le contenu sur la ligne suivante et le @ fermant
            from docx.oxml.ns import qn
            from copy import deepcopy
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            # Paragraphe contenu
            new_p = deepcopy(p._element)
            new_p.find(qn('w:r')).find(qn('w:t')).text = content
            parent.insert(idx + 1, new_p)
            # Paragraphe @ fermant
            close_p = deepcopy(p._element)
            close_p.find(qn('w:r')).find(qn('w:t')).text = "@"
            parent.insert(idx + 2, close_p)
        else:
            p.text = "@dires"
            from docx.oxml.ns import qn
            from copy import deepcopy
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            new_p = deepcopy(p._element)
            new_p.find(qn('w:r')).find(qn('w:t')).text = "(aucune information disponible)"
            parent.insert(idx + 1, new_p)
            close_p = deepcopy(p._element)
            close_p.find(qn('w:r')).find(qn('w:t')).text = "@"
            parent.insert(idx + 2, close_p)
    
    # Remplacer @Analyse par le contenu
    elif text == "@Analyse" and current_pe_section in sections_data:
        analyses = sections_data[current_pe_section].get("analyse", [])
        if analyses:
            content = " ".join(analyses)
            p.text = "@analyse"
            from docx.oxml.ns import qn
            from copy import deepcopy
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            new_p = deepcopy(p._element)
            new_p.find(qn('w:r')).find(qn('w:t')).text = content
            parent.insert(idx + 1, new_p)
            close_p = deepcopy(p._element)
            close_p.find(qn('w:r')).find(qn('w:t')).text = "@"
            parent.insert(idx + 2, close_p)
        else:
            p.text = "@analyse"
            from docx.oxml.ns import qn
            from copy import deepcopy
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            new_p = deepcopy(p._element)
            new_p.find(qn('w:r')).find(qn('w:t')).text = "(aucune information disponible)"
            parent.insert(idx + 1, new_p)
            close_p = deepcopy(p._element)
            close_p.find(qn('w:r')).find(qn('w:t')).text = "@"
            parent.insert(idx + 2, close_p)
    
    i += 1

# Sauvegarder
pea.save("travail/pea.docx")
print("PEA généré : travail/pea.docx")
print(f"Sections trouvées dans le rapport : {list(sections_data.keys())}")
print(f"Nombre de sections avec données : {len(sections_data)}")
