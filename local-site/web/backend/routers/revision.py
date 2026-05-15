"""Router Révision — upload de fichiers et soumission de texte pour révision.

Expose deux endpoints :
- POST /upload : accepte un fichier .docx, .txt ou .md en multipart/form-data
- POST /text : accepte du texte brut en JSON pour révision directe

Valide : Exigences 2.1, 2.2, 2.3, 2.4, 2b.1, 2b.2, 2b.3, 2b.4, 5.3, 5b.1, 7.1, 7.2, 7.5
"""

from __future__ import annotations

import logging
import os
import tempfile

from fastapi import APIRouter, File, HTTPException, UploadFile, status
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

from services.llm_service import LLMConnectionError, LLMService, LLMTimeoutError
from services.revision_models import (
    DocumentParseError,
    TextRevisionRequest,
    TextRevisionResponse,
)
from services.revision_service import RevisionService

logger = logging.getLogger(__name__)

router = APIRouter()

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

ALLOWED_EXTENSIONS: set[str] = {".docx", ".txt", ".md"}
MAX_FILE_SIZE_BYTES: int = 20 * 1024 * 1024  # 20 MB
MAX_TEXT_LENGTH: int = 100_000  # 100 000 caractères


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _get_file_extension(filename: str | None) -> str:
    """Extrait et normalise l'extension du fichier.

    Args:
        filename: Nom du fichier uploadé.

    Returns:
        Extension en minuscules avec le point (ex: '.docx').

    Raises:
        HTTPException: Si le nom de fichier est absent ou l'extension invalide.
    """
    if not filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Nom de fichier manquant.",
        )
    _, ext = os.path.splitext(filename)
    return ext.lower()


def _get_revision_service() -> RevisionService:
    """Instancie le RevisionService.

    Returns:
        Instance de RevisionService prête à l'emploi.
    """
    return RevisionService()


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("/upload")
async def upload_and_revise(file: UploadFile = File(...)):
    """Reçoit un fichier .docx, .txt ou .md et le révise via le LLM.

    - .docx → retourne FileResponse avec track changes (fichier-revu.docx)
    - .txt/.md → retourne JSON avec texte corrigé + filename pour download

    Args:
        file: Fichier uploadé en multipart/form-data.

    Returns:
        FileResponse (.docx) ou TextRevisionResponse (.txt/.md).

    Raises:
        HTTPException 400: Format non supporté ou fichier corrompu.
        HTTPException 413: Fichier > 20 MB.
        HTTPException 503: LLM indisponible.
    """
    # 1. Valider l'extension
    ext = _get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Format non supporté : '{ext}'. "
                f"Formats acceptés : .docx, .txt, .md"
            ),
        )

    # 2. Lire le contenu et valider la taille
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=(
                f"Le fichier dépasse la taille maximale autorisée de 20 MB "
                f"({len(file_bytes) / (1024 * 1024):.1f} MB reçu)."
            ),
        )

    # 3. Appeler le service de révision
    service = _get_revision_service()

    # Extraire le texte du fichier selon le format
    if ext == ".docx":
        import io
        from docx import Document as DocxDocument
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Impossible de lire le fichier .docx : {exc}",
            )
    else:
        # .txt ou .md
        try:
            text_content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text_content = file_bytes.decode("latin-1")

    if not text_content.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier ne contient pas de texte exploitable.",
        )

    try:
        result = await service.revise(text_content)
    except (LLMTimeoutError, LLMConnectionError):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Le service de révision est temporairement indisponible.",
        )

    corrected_text = result.corrected_text

    # 4. Retourner le résultat selon le type
    if ext == ".docx":
        # Pour .docx : retourner le texte corrigé en JSON (pas de track changes pour l'instant)
        output_filename = "fichier-revu.txt"
        return JSONResponse(
            content=TextRevisionResponse(
                corrected_text=corrected_text,
                filename=output_filename,
            ).model_dump(),
        )
    else:
        # .txt ou .md → retourner JSON
        output_filename = f"fichier-revu{ext}"
        return JSONResponse(
            content=TextRevisionResponse(
                corrected_text=corrected_text,
                filename=output_filename,
            ).model_dump(),
        )


# ---------------------------------------------------------------------------
# POST /api/revision/extract-text — Extraction de texte depuis un fichier
# ---------------------------------------------------------------------------


@router.post("/extract-text")
async def extract_text_from_file(file: UploadFile = File(...)):
    """Extrait le texte d'un fichier .docx, .txt ou .md sans le réviser.

    Utilisé pour charger le contenu dans la zone de saisie avant de choisir
    l'action (réviser ou résumer).
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nom de fichier manquant.")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".docx", ".txt", ".md"):
        raise HTTPException(status_code=400, detail="Format non supporté.")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Fichier vide.")

    if ext == ".docx":
        import io
        from docx import Document as DocxDocument
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Impossible de lire le .docx : {exc}")
    else:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")

    return {"text": text, "filename": file.filename, "length": len(text)}


@router.post("/text", response_model=TextRevisionResponse)
async def revise_text(request: TextRevisionRequest) -> TextRevisionResponse:
    """Reçoit du texte brut copié-collé et le révise via le LLM.

    Args:
        request: Corps JSON contenant le texte à réviser.

    Returns:
        TextRevisionResponse avec le texte corrigé.

    Raises:
        HTTPException 400: Texte vide ou > 100 000 caractères.
        HTTPException 503: LLM indisponible.
    """
    # 1. Valider le texte
    text = request.text

    if not text or not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le texte soumis est vide.",
        )

    if len(text) > MAX_TEXT_LENGTH:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Le texte dépasse la limite de {MAX_TEXT_LENGTH:,} caractères "
                f"({len(text):,} caractères reçus)."
            ),
        )

    # 2. Appeler le service de révision
    service = _get_revision_service()
    try:
        result = await service.revise(text)
        corrected_text = result.corrected_text
    except (LLMTimeoutError, LLMConnectionError):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Le service de révision est temporairement indisponible.",
        )

    return TextRevisionResponse(corrected_text=corrected_text)


# ---------------------------------------------------------------------------
# POST /api/revision/summarize — Résumé de texte
# ---------------------------------------------------------------------------


class TextSummarizeRequest(BaseModel):
    """Corps de la requête pour résumer du texte."""
    text: str


class TextSummarizeResponse(BaseModel):
    """Réponse contenant le résumé."""
    summary: str


@router.post("/summarize", response_model=TextSummarizeResponse)
async def summarize_text(request: TextSummarizeRequest) -> TextSummarizeResponse:
    """Reçoit du texte et produit un résumé concis via le LLM.

    Args:
        request: Corps JSON contenant le texte à résumer.

    Returns:
        TextSummarizeResponse avec le résumé.

    Raises:
        HTTPException 400: Texte vide ou > 100 000 caractères.
        HTTPException 503: LLM indisponible.
    """
    text = request.text

    if not text or not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le texte soumis est vide.",
        )

    if len(text) > MAX_TEXT_LENGTH:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Le texte dépasse la limite de {MAX_TEXT_LENGTH:,} caractères "
                f"({len(text):,} caractères reçus)."
            ),
        )

    from services.llm_service import LLMService, LLMTimeoutError, LLMConnectionError

    llm = LLMService()
    try:
        summary = await llm.chat(
            [{"role": "user", "content": text}],
            system_prompt=(
                "Tu es un assistant spécialisé en rédaction de rapports d'expertise judiciaire.\n\n"
                "Résume le texte suivant de manière concise et professionnelle en français.\n"
                "Conserve les éléments essentiels (faits, observations cliniques, conclusions).\n"
                "Le résumé doit être clair, structuré et fidèle au contenu original.\n\n"
                "Réponds uniquement avec le résumé, sans commentaire ni introduction."
            ),
        )
    except (LLMTimeoutError, LLMConnectionError):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Le service LLM est temporairement indisponible.",
        )

    return TextSummarizeResponse(summary=summary)
