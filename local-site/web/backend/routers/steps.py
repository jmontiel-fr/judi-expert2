"""Router des étapes du workflow d'expertise (Step1–Step5).

Valide : Exigences 6.1, 6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 7.1, 7.2, 7.3, 7.4,
         9.1, 9.2, 9.3, 9.4, 9.5, 9.6
"""

from __future__ import annotations

import logging
import os

import httpx
from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, UploadFile, status
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from fastapi.responses import FileResponse

from database import get_db
from models.local_config import LocalConfig
from models.step import Step
from models.step_file import StepFile
from routers.auth import get_current_user
from services.llm_service import LLMService
from services.rag_service import RAGService
from services.workflow_engine import workflow_engine

logger = logging.getLogger(__name__)


def _save_markdown_as_docx(content: str, path: str, title: str) -> None:
    """Convert Markdown-like LLM output to a proper .docx file.

    Parses headings (##, ###) and paragraphs from the LLM text output
    and creates a formatted Word document.

    Args:
        content: Markdown/text content from LLM.
        path: Output file path (.docx).
        title: Document title for the first heading.
    """
    doc = DocxDocument()
    doc.add_heading(title, level=0)

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(path)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OCR_HOST: str = os.environ.get("OCR_HOST", "http://judi-ocr:8001")

# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------


class MarkdownResponse(BaseModel):
    markdown: str


class MarkdownUpdateRequest(BaseModel):
    content: str


class MarkdownUpdateResponse(BaseModel):
    message: str


class ExtractResponse(BaseModel):
    markdown: str
    pdf_path: str
    md_path: str


class Step0ValidateResponse(BaseModel):
    message: str


class QmecResponse(BaseModel):
    qmec: str


class Step1ValidateResponse(BaseModel):
    message: str


class Step2UploadResponse(BaseModel):
    message: str
    filenames: list[str]


class Step2ValidateResponse(BaseModel):
    message: str


class Step3ExecuteResponse(BaseModel):
    message: str
    filenames: list[str]


class Step3ValidateResponse(BaseModel):
    message: str


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _get_dossier_name(dossier_id: int, db: AsyncSession) -> str:
    """Récupère le nom du dossier par son ID."""
    from models.dossier import Dossier as DossierModel
    result = await db.execute(select(DossierModel).where(DossierModel.id == dossier_id))
    dossier = result.scalar_one_or_none()
    if dossier is None:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")
    return dossier.nom


def _step_in(dossier_name: str, step_number: int) -> str:
    """Retourne le chemin du sous-dossier d'entrée d'un step."""
    from services.file_paths import step_in_dir
    return step_in_dir(dossier_name, step_number)


def _step_out(dossier_name: str, step_number: int) -> str:
    """Retourne le chemin du sous-dossier de sortie d'un step."""
    from services.file_paths import step_out_dir
    return step_out_dir(dossier_name, step_number)


def _dossier_dir(dossier_name: str) -> str:
    """Retourne le chemin du répertoire racine d'un dossier."""
    from services.file_paths import dossier_root
    return dossier_root(dossier_name)


async def _get_step(
    dossier_id: int, step_number: int, db: AsyncSession
) -> Step:
    """Charge l'étape demandée avec ses fichiers ou lève 404."""
    result = await db.execute(
        select(Step)
        .options(selectinload(Step.files))
        .where(Step.dossier_id == dossier_id, Step.step_number == step_number)
    )
    step = result.scalar_one_or_none()
    if step is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Étape {step_number} non trouvée pour le dossier {dossier_id}",
        )
    return step


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------

router = APIRouter()


# ---- Step1 : POST /api/dossiers/{id}/step1/upload -------------------------


class UploadResponse(BaseModel):
    message: str
    filename: str
    file_size: int


@router.post(
    "/{dossier_id}/step1/upload",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step1_upload(
    dossier_id: int,
    file: UploadFile,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload le PDF de demande dans step1/in/ sans lancer de traitement.

    Le fichier conserve son nom d'origine (ex: requisition.pdf, ordonnance.pdf).
    """

    # Vérifier que l'étape n'est pas verrouillée
    await workflow_engine.require_step_not_validated(dossier_id, 1, db)

    # Vérifier que le fichier est un PDF
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Seuls les fichiers PDF sont acceptés",
        )

    # Lire le contenu
    pdf_content = await file.read()
    if not pdf_content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier PDF est vide",
        )

    # Conserver le nom d'origine du fichier
    original_filename = file.filename

    # Résoudre le nom du dossier et sauvegarder
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 1)
    os.makedirs(in_dir, exist_ok=True)

    pdf_path = os.path.join(in_dir, original_filename)
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)

    # Créer le StepFile en base
    step = await _get_step(dossier_id, 1, db)

    # Supprimer l'ancien fichier principal s'il existe (ré-upload)
    for old_file in list(step.files):
        if old_file.file_type == "pdf_scan":
            await db.delete(old_file)
    await db.flush()

    step_file = StepFile(
        step_id=step.id,
        filename=original_filename,
        file_path=pdf_path,
        file_type="pdf_scan",
        file_size=len(pdf_content),
    )
    db.add(step_file)
    await db.commit()

    return UploadResponse(
        message="Fichier importé avec succès",
        filename=original_filename,
        file_size=len(pdf_content),
    )


# ---- Step1 : POST /api/dossiers/{id}/step1/execute -----------------------

@router.post(
    "/{dossier_id}/step1/execute",
    response_model=ExtractResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step1_execute(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Lance l'OCR + structuration LLM sur les fichiers déjà uploadés dans step1/in/.

    Produit : demande.md, placeholders.csv dans step1/out/.
    """

    # 0. Vérifier que l'étape n'est pas verrouillée
    await workflow_engine.require_step_not_validated(dossier_id, 1, db)

    # 0b. Vérifier que le step n'est pas déjà en cours (éviter re-lancement sur refresh)
    step = await _get_step(dossier_id, 1, db)
    if step.statut == "en_cours":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Le traitement est déjà en cours",
        )

    # 0c. Trouver le fichier PDF uploadé (quel que soit son nom)
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 1)
    out_dir = _step_out(dossier_name, 1)

    # Chercher le premier .pdf dans step1/in/
    pdf_path = None
    if os.path.isdir(in_dir):
        for fname in os.listdir(in_dir):
            if fname.lower().endswith(".pdf"):
                pdf_path = os.path.join(in_dir, fname)
                break

    if not pdf_path or not os.path.isfile(pdf_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Aucun fichier PDF uploadé — importez d'abord le PDF",
        )

    # 1. Marquer le step comme "en_cours" avec progression
    await workflow_engine.start_step(dossier_id, 1, db)
    step = await _get_step(dossier_id, 1, db)
    step.progress_current = 1
    step.progress_total = 5
    step.progress_message = "OCR (extraction du texte)"
    await db.commit()

    # 2. Lire le PDF
    with open(pdf_path, "rb") as f:
        pdf_content = f.read()

    # 3. Appeler le service OCR
    logger.info("[Step1] Dossier %d — Phase 1/3 : OCR en cours…", dossier_id)
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                f"{OCR_HOST}/api/ocr/extract",
                files={"file": ("ordonnance.pdf", pdf_content, "application/pdf")},
            )
            resp.raise_for_status()
            ocr_data = resp.json()
            texte_brut = ocr_data.get("text", "")
    except httpx.ConnectError:
        await workflow_engine.fail_step(dossier_id, 1, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Service OCR indisponible — vérifiez que le conteneur judi-ocr est démarré",
        )
    except httpx.TimeoutException:
        await workflow_engine.fail_step(dossier_id, 1, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Service OCR indisponible — délai de connexion dépassé",
        )
    except httpx.HTTPStatusError as exc:
        logger.error("Erreur OCR HTTP %s", exc.response.status_code)
        await workflow_engine.fail_step(dossier_id, 1, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Erreur du service OCR lors de l'extraction",
        )

    if not texte_brut.strip():
        await workflow_engine.fail_step(dossier_id, 1, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Le PDF ne contient pas de texte exploitable",
        )

    # 3b. Vérifier si l'opération a été annulée entre OCR et LLM
    step_check = await _get_step(dossier_id, 1, db)
    if step_check.statut != "en_cours":
        logger.info("Step 1 annulé après OCR — abandon du traitement")
        return ExtractResponse(markdown="", pdf_path=pdf_path, md_path="")

    # 4. Appeler le LLM pour structurer en Markdown
    logger.info("[Step1] Dossier %d — Phase 2/5 : Structuration ordonnance (LLM)…", dossier_id)
    step = await _get_step(dossier_id, 1, db)
    step.progress_current = 2
    step.progress_message = "Structuration en Markdown ⏳ LLM"
    await db.commit()
    llm = LLMService()
    try:
        markdown = await llm.structurer_markdown(texte_brut)
    except Exception as exc:
        logger.error("Erreur LLM lors de la structuration Markdown : %s", exc)
        await workflow_engine.fail_step(dossier_id, 1, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Service LLM indisponible — essayez de redémarrer le conteneur judi-llm",
        )

    # 4b. Vérifier si l'opération a été annulée pendant le LLM
    await db.refresh(step_check)
    if step_check.statut != "en_cours":
        logger.info("Step 1 annulé après LLM — abandon du traitement")
        return ExtractResponse(markdown="", pdf_path=pdf_path, md_path="")

    # 4b. Nettoyer le markdown (supprimer les blocs ``` ajoutés par le LLM)
    markdown = markdown.strip()
    if markdown.startswith("```markdown"):
        markdown = markdown[len("```markdown"):].strip()
    elif markdown.startswith("```md"):
        markdown = markdown[len("```md"):].strip()
    elif markdown.startswith("```"):
        markdown = markdown[3:].strip()
    if markdown.endswith("```"):
        markdown = markdown[:-3].strip()

    # 5. Sauvegarder les fichiers de sortie
    step = await _get_step(dossier_id, 1, db)
    step.progress_current = 5
    step.progress_message = "Sauvegarde des fichiers"
    await db.commit()

    os.makedirs(out_dir, exist_ok=True)

    md_path = os.path.join(out_dir, "demande.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(markdown)

    # 5b. Vérifier annulation avant extraction des questions
    step_check = await _get_step(dossier_id, 1, db)
    if step_check.statut != "en_cours":
        logger.info("Step 1 annulé avant extraction questions — abandon")
        return ExtractResponse(markdown=markdown, pdf_path=pdf_path, md_path=md_path)

    # 6. Extraire les questions du tribunal
    logger.info("[Step1] Dossier %d — Phase 3/5 : Extraction des questions (LLM)…", dossier_id)
    step = await _get_step(dossier_id, 1, db)
    step.progress_current = 3
    step.progress_message = "Extraction des questions ⏳ LLM"
    await db.commit()
    try:
        questions_md = await llm.extraire_questions(markdown)
    except Exception as exc:
        logger.warning("Extraction questions échouée : %s — on continue sans", exc)
        questions_md = ""

    # Nettoyer le markdown des questions
    if questions_md:
        questions_md = questions_md.strip()
        if questions_md.startswith("```"):
            questions_md = questions_md.split("\n", 1)[-1] if "\n" in questions_md else ""
        if questions_md.endswith("```"):
            questions_md = questions_md[:-3].strip()

    # 6b. Vérifier annulation avant extraction des placeholders
    await db.refresh(step_check)
    if step_check.statut != "en_cours":
        logger.info("Step 1 annulé avant extraction placeholders — abandon")
        return ExtractResponse(markdown=markdown, pdf_path=pdf_path, md_path=md_path)

    # 7. Extraire les placeholders
    logger.info("[Step1] Dossier %d — Phase 4/5 : Extraction des placeholders (LLM)…", dossier_id)
    step = await _get_step(dossier_id, 1, db)
    step.progress_current = 4
    step.progress_message = "Extraction des placeholders ⏳ LLM"
    await db.commit()
    try:
        placeholders_csv = await llm.extraire_placeholders(markdown)
    except Exception as exc:
        logger.warning("Extraction placeholders échouée : %s — on continue sans", exc)
        placeholders_csv = ""

    # Nettoyer le CSV
    if placeholders_csv:
        placeholders_csv = placeholders_csv.strip()
        if placeholders_csv.startswith("```"):
            placeholders_csv = placeholders_csv.split("\n", 1)[-1] if "\n" in placeholders_csv else ""
        if placeholders_csv.endswith("```"):
            placeholders_csv = placeholders_csv[:-3].strip()

    # Réparer les lignes collées (le LLM oublie parfois le retour à la ligne)
    # Pattern : "valeurnom_placeholder;" → séparer en deux lignes
    if placeholders_csv:
        import re as _re_csv
        # Détecter "texte_sans_newlineNOM_CONNU;valeur" et insérer un \n avant
        _KNOWN_KEYS = (
            "titre_expertise", "objet_mission", "date_mission",
            "reference_dossier", "pieces_auxiliaires",
            "requerant_prenom", "requerant_nom", "requerant_titre", "requerant_ville",
            "nom_tribunal", "ville_tribunal", "nom_magistrat",
            "genre_pex", "nom_pex", "prenom_pex", "age_pex",
            "date_naissance_pex", "ville_naissance_pex", "CP_ville_naissance_pex",
            "genre_expert", "nom_expert", "prenom_expert", "titre_expert",
            "date_rapport", "question_",
        )
        for known_key in _KNOWN_KEYS:
            # Remplacer "texteKEY;" par "texte\nKEY;" quand KEY n'est pas en début de ligne
            pattern = _re_csv.compile(r"([^\n])(" + _re_csv.escape(known_key) + r"[^;]*;)")
            placeholders_csv = pattern.sub(r"\1\n\2", placeholders_csv)

    # Filtrer : ne conserver que les placeholders de réquisition standards
    # (définis dans glossaire-workflow) et uniquement s'ils sont valorisés
    _PLACEHOLDERS_REQUISITION_STANDARDS = {
        "titre_expertise", "objet_mission", "date_mission",
        "reference_dossier",
        "requerant_prenom", "requerant_nom", "requerant_titre", "requerant_ville",
        "nom_tribunal", "ville_tribunal", "nom_magistrat",
        "genre_pex", "nom_pex", "prenom_pex",
        "date_naissance_pex", "ville_naissance_pex",
        "genre_expert", "nom_expert", "prenom_expert", "titre_expert",
    }

    filtered_lines = ["nom_placeholder;valeur"]
    if placeholders_csv.strip():
        for line in placeholders_csv.strip().split("\n"):
            line = line.strip()
            if not line or line.startswith("nom_placeholder"):
                continue
            parts = line.split(";", 1)
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()
                # Nettoyer les retours à la ligne dans les valeurs
                value = value.replace("\n", " ").replace("\r", " ").strip()
                # Ignorer les valeurs "..." (placeholder non rempli par le LLM)
                if value == "..." or value == "…":
                    continue
                # Ne garder que les standards ET valorisés
                if key in _PLACEHOLDERS_REQUISITION_STANDARDS and value:
                    filtered_lines.append(f"{key};{value}")

    placeholders_csv = "\n".join(filtered_lines) if len(filtered_lines) > 1 else ""

    placeholders_path = ""
    if placeholders_csv.strip():
        placeholders_path = os.path.join(out_dir, "placeholders.csv")
        with open(placeholders_path, "w", encoding="utf-8") as f:
            f.write(placeholders_csv)

    # 7b. Sauvegarder questions.md et ajouter les questions dans placeholders.csv
    questions_path = ""
    if questions_md.strip():
        questions_path = os.path.join(out_dir, "questions.md")
        with open(questions_path, "w", encoding="utf-8") as f:
            f.write(questions_md)

        # Parser les questions depuis le markdown pour les ajouter au CSV
        import re as _re
        question_lines = []
        # Format attendu : ## Q1\nTexte... ou ## Q1 :\nTexte...
        q_pattern = _re.compile(r"##\s*Q(\d+)[^\n]*\n(.*?)(?=\n##\s*Q\d|\Z)", _re.DOTALL)
        matches = list(q_pattern.finditer(questions_md))

        if not matches:
            # Fallback : format tirets "- Texte"
            q_lines = [l.strip().lstrip("- ").strip() for l in questions_md.strip().split("\n")
                       if l.strip() and not l.strip().startswith("#")]
            for i, q_text in enumerate(q_lines, 1):
                if q_text:
                    q_text = q_text.replace("\n", " ").strip()
                    question_lines.append(f"question_{i};{q_text}")
        else:
            for match in matches:
                q_num = match.group(1)
                q_text = match.group(2).strip().replace("\n", " ")
                q_text = _re.sub(r"\s+", " ", q_text).strip()
                if q_text:
                    question_lines.append(f"question_{q_num};{q_text}")

        if question_lines:
            if not placeholders_path:
                placeholders_path = os.path.join(out_dir, "placeholders.csv")
                with open(placeholders_path, "w", encoding="utf-8") as f:
                    f.write("nom_placeholder;valeur\n")
                    f.write("\n".join(question_lines) + "\n")
            else:
                with open(placeholders_path, "a", encoding="utf-8") as f:
                    f.write("\n".join(question_lines) + "\n")

    # 8. Créer les entrées StepFile de sortie en base
    step = await _get_step(dossier_id, 1, db)

    # Supprimer les anciens fichiers de sortie (ré-exécution)
    for old_file in list(step.files):
        if old_file.file_type in ("markdown", "questions", "placeholders"):
            await db.delete(old_file)
    await db.flush()

    md_size = len(markdown.encode("utf-8"))
    step_file_md = StepFile(
        step_id=step.id,
        filename="demande.md",
        file_path=md_path,
        file_type="markdown",
        file_size=md_size,
    )
    db.add(step_file_md)

    if questions_path:
        db.add(StepFile(
            step_id=step.id,
            filename="questions.md",
            file_path=questions_path,
            file_type="questions",
            file_size=len(questions_md.encode("utf-8")),
        ))

    if placeholders_path:
        db.add(StepFile(
            step_id=step.id,
            filename="placeholders.csv",
            file_path=placeholders_path,
            file_type="placeholders",
            file_size=len(placeholders_csv.encode("utf-8")),
        ))

    # 9. Marquer step1 comme "fait"
    logger.info("[Step1] Dossier %d — Terminé. Fichiers générés : demande.md%s",
                dossier_id,
                ", placeholders.csv" if placeholders_path else "")
    await workflow_engine.execute_step(dossier_id, 1, db)
    await db.commit()

    return ExtractResponse(
        markdown=markdown,
        pdf_path=pdf_path,
        md_path=md_path,
    )


# ---- Step1 : POST /api/dossiers/{id}/step1/extract (legacy — redirige) ---

@router.post(
    "/{dossier_id}/step1/extract",
    response_model=ExtractResponse,
    status_code=status.HTTP_201_CREATED,
    deprecated=True,
)
async def step0_extract(
    dossier_id: int,
    file: UploadFile,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """[LEGACY] Upload + extraction en un seul appel. Utiliser /upload puis /execute."""
    # Upload d'abord
    await step1_upload.__wrapped__(dossier_id, file, _user, db) if hasattr(step1_upload, '__wrapped__') else None

    # Sauvegarder le fichier manuellement (fallback si __wrapped__ n'existe pas)
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 1)
    os.makedirs(in_dir, exist_ok=True)

    pdf_content = await file.read()
    if not pdf_content:
        # Le fichier a déjà été lu par step1_upload, relire depuis le disque
        pdf_path = os.path.join(in_dir, "ordonnance.pdf")
        if os.path.isfile(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_content = f.read()

    if not pdf_content:
        raise HTTPException(status_code=400, detail="Fichier vide")

    # Puis exécuter
    return await step1_execute(dossier_id, _user, db)


# ---- Step1 : GET /api/dossiers/{id}/step1/markdown -----------------------

@router.get(
    "/{dossier_id}/step1/markdown",
    response_model=MarkdownResponse,
)
async def step0_get_markdown(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Retourne le contenu du fichier Markdown généré pour step0."""

    # Vérifier l'accès à l'étape
    await workflow_engine.require_step_access(dossier_id, 1, db)

    dossier_name = await _get_dossier_name(dossier_id, db)
    md_path = os.path.join(_step_out(dossier_name, 1), "ordonnance.md")
    if not os.path.isfile(md_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Fichier Markdown non trouvé — lancez d'abord l'extraction",
        )

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    return MarkdownResponse(markdown=content)


# ---- Step1 : PUT /api/dossiers/{id}/step1/markdown -----------------------

@router.put(
    "/{dossier_id}/step1/markdown",
    response_model=MarkdownUpdateResponse,
)
async def step0_update_markdown(
    dossier_id: int,
    body: MarkdownUpdateRequest,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Met à jour le fichier Markdown (édition manuelle par l'expert)."""

    # Vérifier que l'étape n'est pas validée (immuable)
    await workflow_engine.require_step_not_validated(dossier_id, 1, db)

    dossier_name = await _get_dossier_name(dossier_id, db)
    md_path = os.path.join(_step_out(dossier_name, 1), "ordonnance.md")
    if not os.path.isfile(md_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Fichier Markdown non trouvé — lancez d'abord l'extraction",
        )

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(body.content)

    # Mettre à jour la taille du fichier dans StepFile
    step = await _get_step(dossier_id, 1, db)
    for sf in step.files:
        if sf.filename == "ordonnance.md":
            sf.file_size = len(body.content.encode("utf-8"))
            break

    await db.commit()

    return MarkdownUpdateResponse(message="Fichier Markdown mis à jour")


# ---- Step1 : POST /api/dossiers/{id}/step1/import-docx -------------------

@router.post(
    "/{dossier_id}/step1/import-docx",
    response_model=MarkdownUpdateResponse,
)
async def step0_import_docx(
    dossier_id: int,
    file: UploadFile,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Importe un .docx modifié par l'expert et met à jour le .md et le .docx."""
    from docx import Document as DocxDocument

    await workflow_engine.require_step_not_validated(dossier_id, 1, db)

    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Seul le format .docx est accepté",
        )

    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier est vide",
        )

    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 1)
    out_dir = _step_out(dossier_name, 1)

    # Sauvegarder le .docx importé (input — uploadé par l'expert)
    docx_path = os.path.join(in_dir, "ordonnance.docx")
    with open(docx_path, "wb") as f:
        f.write(content)

    # Extraire le texte du .docx pour mettre à jour le .md (output)
    import io
    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    markdown = "\n\n".join(paragraphs)

    md_path = os.path.join(out_dir, "ordonnance.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(markdown)

    # Mettre à jour les tailles dans StepFile
    step = await _get_step(dossier_id, 1, db)
    for sf in step.files:
        if sf.filename == "ordonnance.md":
            sf.file_size = len(markdown.encode("utf-8"))
        elif sf.filename == "ordonnance.docx":
            sf.file_size = len(content)

    await db.commit()

    return MarkdownUpdateResponse(message="Document importé avec succès")


# ---- Step1 : POST /api/dossiers/{id}/step1/complementary -----------------

# Types de pièces complémentaires
VALID_DOC_TYPES = ("rapport", "plainte", "autre")
# Formats acceptés et ceux qui passent à l'OCR
VALID_DOC_FORMATS = ("pdf", "scan", "image", "csv", "xlsx")
OCR_FORMATS = ("pdf", "scan")


class ComplementaryResponse(BaseModel):
    message: str
    filename: str
    ocr_applied: bool
    output_file: str | None = None


@router.post(
    "/{dossier_id}/step1/complementary",
    response_model=ComplementaryResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step1_upload_complementary(
    dossier_id: int,
    file: UploadFile,
    label: str = "",
    doc_type: str = "autre",
    doc_format: str = "pdf",
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload une pièce complémentaire au dossier (Step 1).

    - type : rapport, plainte, autre
    - format : pdf, scan, image, csv, xlsx
    - Seuls les formats pdf et scan passent à l'OCR → produisent un .md
    - Les formats csv et xlsx sont stockés tels quels
    """
    # Validation
    if doc_type not in VALID_DOC_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Type invalide. Valeurs acceptées : {', '.join(VALID_DOC_TYPES)}",
        )
    if doc_format not in VALID_DOC_FORMATS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Format invalide. Valeurs acceptées : {', '.join(VALID_DOC_FORMATS)}",
        )

    # Vérifier que l'étape n'est pas verrouillée
    await workflow_engine.require_step_not_validated(dossier_id, 1, db)

    # Lire le fichier
    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier est vide",
        )

    # Créer le répertoire
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 1)
    out_dir = _step_out(dossier_name, 1)
    os.makedirs(in_dir, exist_ok=True)
    os.makedirs(out_dir, exist_ok=True)

    # Sauvegarder le fichier original (input — uploadé par l'expert)
    safe_filename = file.filename or f"piece_{doc_type}.{doc_format}"
    file_path = os.path.join(in_dir, safe_filename)
    with open(file_path, "wb") as f:
        f.write(content)

    # OCR si format pdf ou scan
    ocr_applied = False
    output_file = None

    if doc_format in OCR_FORMATS:
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                resp = await client.post(
                    f"{OCR_HOST}/api/ocr/extract",
                    files={"file": (safe_filename, content, "application/pdf")},
                )
                resp.raise_for_status()
                ocr_data = resp.json()
                texte_brut = ocr_data.get("text", "")

            if texte_brut.strip():
                # Sauvegarder le résultat OCR en .md (output — généré par OCR)
                md_filename = os.path.splitext(safe_filename)[0] + ".md"
                md_path = os.path.join(out_dir, md_filename)
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(texte_brut)
                ocr_applied = True
                output_file = md_filename

                # Créer le StepFile pour le .md
                step = await _get_step(dossier_id, 1, db)
                md_step_file = StepFile(
                    step_id=step.id,
                    filename=md_filename,
                    file_path=md_path,
                    file_type="complementary_ocr",
                    file_size=len(texte_brut.encode("utf-8")),
                    doc_type=doc_type,
                    doc_format=doc_format,
                )
                db.add(md_step_file)
        except (httpx.ConnectError, httpx.TimeoutException, httpx.HTTPStatusError) as e:
            logger.warning("OCR indisponible pour pièce complémentaire : %s", e)

    # Créer le StepFile pour le fichier original
    step = await _get_step(dossier_id, 1, db)
    step_file = StepFile(
        step_id=step.id,
        filename=safe_filename,
        file_path=file_path,
        file_type="complementary",
        file_size=len(content),
        doc_type=doc_type,
        doc_format=doc_format,
    )
    db.add(step_file)
    await db.commit()

    return ComplementaryResponse(
        message=f"Pièce complémentaire '{label or safe_filename}' ajoutée.",
        filename=safe_filename,
        ocr_applied=ocr_applied,
        output_file=output_file,
    )


# ---- Step1 : POST /api/dossiers/{id}/step1/validate ---------------------

@router.post(
    "/{dossier_id}/step1/validate",
    response_model=Step0ValidateResponse,
)
async def step1_validate(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Valide l'étape 1 (extraction) — passage à "validé"."""

    await workflow_engine.validate_step(dossier_id, 1, db)
    await db.commit()

    return Step0ValidateResponse(message="Step 1 validé avec succès")


# ---- Step2 : POST /api/dossiers/{id}/step2/execute -----------------------

@router.post(
    "/{dossier_id}/step2/execute",
    response_model=QmecResponse,
)
async def step2_execute(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Extrait le PE (plan d'entretien) depuis le TRE.

    Nouveau flux TRE-centré :
    1. Résoudre le TRE (personnalisé ou par défaut du domaine)
    2. Valider la structure du TRE (présence @debut_tpe@)
    3. Charger les questions depuis placeholders.csv du Step 1
    4. Extraire le PE : portion @debut_tpe@ → fin + questions en conclusion
    5. Sauvegarder pe.docx

    Valide : Exigences 7.1, 7.2 (Requirement 2)
    """
    import shutil

    from services.tre_parser import TREParser
    from services.file_paths import tre_path as resolve_tre_path

    # 1. Vérifier l'accès au step
    await workflow_engine.require_step_access(dossier_id, 2, db)

    # Résoudre le nom du dossier pour les chemins
    dossier_name = await _get_dossier_name(dossier_id, db)

    # 2. Vérifier que le step n'est pas déjà en cours
    step = await _get_step(dossier_id, 2, db)
    if step.statut == "en_cours":
        return QmecResponse(qmec="")

    # 3. Récupérer le domaine
    result = await db.execute(select(LocalConfig).limit(1))
    config = result.scalar_one_or_none()
    if config is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Configuration initiale non effectuée",
        )
    domaine = config.domaine

    # 4. Résoudre le TRE
    tre_file = resolve_tre_path(dossier_name, domaine)
    if tre_file is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Aucun TRE trouvé — uploadez un TRE personnalisé ou placez un fichier tre.docx dans le corpus du domaine",
        )

    # 5. Marquer le step comme "en_cours"
    step = await workflow_engine.start_step(dossier_id, 2, db)
    step.progress_current = 1
    step.progress_total = 3
    step.progress_message = "Validation syntaxique du TRE"
    await db.commit()

    # 6. Copier le TRE dans step2/in/ (figer la version)
    step2_in = _step_in(dossier_name, 2)
    os.makedirs(step2_in, exist_ok=True)
    tre_local = os.path.join(step2_in, "tre.docx")
    if tre_file != tre_local:
        shutil.copy2(tre_file, tre_local)

    # 7. Parser et valider le TRE
    parser = TREParser()
    try:
        parse_result = parser.parse(tre_local)
    except Exception as exc:
        logger.error("[Step2] Erreur parsing TRE : %s", exc)
        await workflow_engine.fail_step(dossier_id, 2, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Erreur lors du parsing du TRE : {exc}",
        )

    errors = parser.validate(parse_result)
    if errors:
        await workflow_engine.fail_step(dossier_id, 2, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"TRE invalide : {'; '.join(errors)}",
        )

    # 8. Vérifier que les placeholders du TRE sont définis dans placeholders.csv
    placeholders_path = os.path.join(_step_out(dossier_name, 1), "placeholders.csv")
    defined_placeholders: set[str] = set()
    if os.path.isfile(placeholders_path):
        with open(placeholders_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("nom_placeholder"):
                    continue
                parts = line.split(";", 1)
                if len(parts) >= 1 and parts[0].strip():
                    defined_placeholders.add(parts[0].strip())

    # Vérifier les placeholders du TRE contre ceux définis
    tre_placeholder_names = {p.name for p in parse_result.placeholders}
    undefined_placeholders = tre_placeholder_names - defined_placeholders
    # Exclure les placeholders qui seront remplis par d'autres sources :
    # - annotations (dires_*, analyse_*, verbatim_*)
    # - placeholders de configuration locale (genre_expert, nom_expert, prenom_expert, titre_expert)
    # - placeholders calculés automatiquement (date_rapport, age_pex)
    # - questions (question_*)
    _AUTO_FILLED = {"genre_expert", "nom_expert", "prenom_expert", "titre_expert",
                    "date_rapport", "age_pex", "pieces_auxiliaires"}
    undefined_placeholders = {
        p for p in undefined_placeholders
        if not any(p.startswith(prefix) for prefix in ("dires_", "analyse_", "verbatim_", "question_"))
        and p not in _AUTO_FILLED
    }

    placeholder_warning = ""
    if undefined_placeholders:
        placeholder_warning = (
            f"Attention : {len(undefined_placeholders)} placeholder(s) du TRE "
            f"non défini(s) dans placeholders.csv : {', '.join(sorted(undefined_placeholders))}. "
            f"Ils seront remplacés par '[nom_placeholder]' au Step 4."
        )
        logger.warning("[Step2] Dossier %d — %s", dossier_id, placeholder_warning)

    # 9. Copier le TRE complet en sortie (l'expert l'annotera directement)
    step = await _get_step(dossier_id, 2, db)
    step.progress_current = 2
    step.progress_message = "Vérification des <<placeholders>> contre placeholders.csv"
    await db.commit()

    # (La vérification des placeholders est déjà faite ci-dessus — étape 2 = juste le message)

    # 10. Sauvegarder le TRE complet en sortie
    step = await _get_step(dossier_id, 2, db)
    step.progress_current = 3
    step.progress_message = "Copie du TRE complet en sortie"
    await db.commit()

    step2_out = _step_out(dossier_name, 2)
    os.makedirs(step2_out, exist_ok=True)

    # Copier le TRE complet (pas d'extraction PE — l'expert annote le TRE directement)
    import shutil as _shutil
    tre_out_path = os.path.join(step2_out, "tre.docx")
    _shutil.copy2(tre_local, tre_out_path)

    # 11. Créer les StepFile en base
    step = await _get_step(dossier_id, 2, db)
    for old_file in list(step.files):
        await db.delete(old_file)
    await db.flush()

    tre_size = os.path.getsize(tre_out_path)
    db.add(StepFile(
        step_id=step.id,
        filename="tre.docx",
        file_path=tre_out_path,
        file_type="plan_entretien_docx",
        file_size=tre_size,
    ))

    # 12. Marquer step2 comme "réalisé"
    await workflow_engine.execute_step(dossier_id, 2, db)
    step = await _get_step(dossier_id, 2, db)
    step.progress_current = None
    step.progress_total = None
    step.progress_message = None
    await db.commit()

    logger.info("[Step2] Dossier %d — TRE validé et copié avec succès", dossier_id)

    return QmecResponse(qmec=placeholder_warning)



# ---- Step2 : POST /api/dossiers/{id}/step2/upload-tre --------------------

@router.post(
    "/{dossier_id}/step2/upload-tre",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step2_upload_tre(
    dossier_id: int,
    file: UploadFile,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload d'un TRE personnalisé pour ce dossier.

    Le TRE uploadé sera utilisé au Step 2 (extraction PE) et au Step 4
    (reconstitution du rapport). Il est stocké dans step2/in/tre.docx.
    """
    # Vérifier l'accès
    await workflow_engine.require_step_access(dossier_id, 2, db)
    await workflow_engine.require_step_not_validated(dossier_id, 2, db)

    # Valider le format
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Seul le format .docx est accepté pour le TRE",
        )

    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier est vide",
        )

    # Valider la structure du TRE
    import io
    from services.tre_parser import TREParser

    parser = TREParser()
    try:
        from docx import Document as DocxDocument
        # Vérifier que c'est un docx valide avec @debut_tpe@
        temp_doc = DocxDocument(io.BytesIO(content))
        found_debut_tpe = False
        for para in temp_doc.paragraphs:
            if parser.DEBUT_TPE_RE.match(para.text):
                found_debut_tpe = True
                break
        if not found_debut_tpe:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Le TRE doit contenir le marqueur @debut_tpe@ pour délimiter le PE",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Fichier .docx invalide : {exc}",
        )

    # Sauvegarder le TRE
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 2)
    os.makedirs(in_dir, exist_ok=True)

    tre_path_local = os.path.join(in_dir, "tre.docx")
    with open(tre_path_local, "wb") as f:
        f.write(content)

    # Créer/mettre à jour le StepFile en base
    step = await _get_step(dossier_id, 2, db)
    for old_file in list(step.files):
        if old_file.filename == "tre.docx":
            await db.delete(old_file)
    await db.flush()

    step_file = StepFile(
        step_id=step.id,
        filename="tre.docx",
        file_path=tre_path_local,
        file_type="tre",
        file_size=len(content),
    )
    db.add(step_file)
    await db.commit()

    return UploadResponse(
        message="TRE personnalisé importé avec succès",
        filename="tre.docx",
        file_size=len(content),
    )


# ---- Step2 : GET /api/dossiers/{id}/step2/download -----------------------

@router.get(
    "/{dossier_id}/step2/download",
)
async def step2_download(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Télécharge le fichier QMEC généré pour step1.

    Valide : Exigence 7.3
    """

    # Vérifier l'accès à l'étape
    await workflow_engine.require_step_access(dossier_id, 2, db)

    dossier_name = await _get_dossier_name(dossier_id, db)
    step2_out = _step_out(dossier_name, 2)

    pe_path = os.path.join(step2_out, "pe.docx")
    if not os.path.isfile(pe_path):
        # Fallback to .md
        pe_path = os.path.join(step2_out, "pe.md")
        if not os.path.isfile(pe_path):
            # Legacy fallback
            pe_path = os.path.join(step2_out, "qmec.docx")
            if not os.path.isfile(pe_path):
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Plan d'Entretien non trouvé — lancez d'abord l'exécution du Step 2",
                )
        return FileResponse(
            path=pe_path,
            filename="pe.md",
            media_type="text/markdown",
        )

    return FileResponse(
        path=pe_path,
        filename="pe.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---- Step2 : POST /api/dossiers/{id}/step2/validate ---------------------

@router.post(
    "/{dossier_id}/step2/validate",
    response_model=Step1ValidateResponse,
)
async def step2_validate(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Valide le Step 2 (verrouillage) et autorise l'accès au Step 3.

    Valide : Exigence 7.4
    """

    await workflow_engine.validate_step(dossier_id, 2, db)
    await db.commit()

    return Step1ValidateResponse(message="Step 2 validé avec succès")


# ---- Step3 : POST /api/dossiers/{id}/step3/skip — Valider sans pièces ------

@router.post(
    "/{dossier_id}/step3/skip",
    response_model=Step1ValidateResponse,
)
async def step3_skip(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Valide le Step 3 sans pièces de consolidation (sans objet).

    Passe directement de 'initial' à 'validé' sans exécution.
    """
    step = await _get_step(dossier_id, 3, db)
    if step.statut == "valide":
        return Step1ValidateResponse(message="Step 3 déjà validé")
    if step.statut not in ("initial", "fait"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="L'étape doit être au statut 'initial' ou 'fait' pour être passée",
        )

    # Vérifier les prérequis (step 2 validé)
    await workflow_engine.require_step_access(dossier_id, 3, db)

    # Passer directement à validé
    from datetime import datetime, UTC
    step.statut = "valide"
    step.executed_at = datetime.now(UTC)
    step.validated_at = datetime.now(UTC)
    step.execution_duration_seconds = 0
    await db.commit()

    return Step1ValidateResponse(message="Step 3 validé (sans objet)")


# ---- Step4 : POST /api/dossiers/{id}/step4/upload — Import PEA uniquement ----

@router.post(
    "/{dossier_id}/step4/upload",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step4_upload(
    dossier_id: int,
    file: UploadFile,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload du PEA/PAA (.docx) sans lancer la génération du pré-rapport."""

    # Vérifier l'accès
    await workflow_engine.require_step_access(dossier_id, 4, db)
    await workflow_engine.require_step_not_validated(dossier_id, 4, db)

    # Valider le format
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Seul le format .docx est accepté",
        )

    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le fichier est vide",
        )

    # Sauvegarder le PEA
    dossier_name = await _get_dossier_name(dossier_id, db)
    in_dir = _step_in(dossier_name, 4)
    os.makedirs(in_dir, exist_ok=True)

    pea_path = os.path.join(in_dir, "pea.docx")
    with open(pea_path, "wb") as f:
        f.write(content)

    # Créer le StepFile en base
    step = await _get_step(dossier_id, 4, db)

    # Supprimer l'ancien PEA s'il existe
    for old_file in list(step.files):
        if old_file.filename == "pea.docx":
            await db.delete(old_file)
    await db.flush()

    step_file = StepFile(
        step_id=step.id,
        filename="pea.docx",
        file_path=pea_path,
        file_type="pea",
        file_size=len(content),
    )
    db.add(step_file)
    await db.commit()

    return UploadResponse(
        message="PEA/PAA importé avec succès",
        filename="pea.docx",
        file_size=len(content),
    )


# ---- Step4 : POST /api/dossiers/{id}/step4/execute ------------------------

@router.post(
    "/{dossier_id}/step4/execute",
    response_model=Step2UploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step4_execute(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Génération du PRE (Pré-Rapport d'Expertise).

    Approche : find/replace dans le .docx du PEA.
    Le PEA est une copie du TRE annotée par l'expert. On travaille directement
    dessus pour préserver styles, table des matières, numérotation, etc.

    Flow :
    1. Validation : parser le PEA, vérifier annotations et placeholders
    2. Reformulation LLM des @dires et @analyse
    3. Résolution des @resume (concaténation + résumé LLM)
    4. Résolution @question, @reference, @cite
    5. Substitution in-place dans le .docx : annotations → texte reformulé,
       <<placeholders>> → valeurs, @remplir → texte après ":"

    Valide : Exigences 3.1, 3.2, 3.3, 3.4, 3.5, 3.6, 3.7, 3.8
    """
    import re as re_module
    import shutil

    from docx import Document as DocxDocument
    from services.tre_parser import TREParser

    # 1. Vérifier l'accès au step4
    await workflow_engine.require_step_access(dossier_id, 4, db)

    # 1b. Marquer le step comme "en_cours"
    await workflow_engine.start_step(dossier_id, 4, db)
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 1
    step.progress_total = 5
    step.progress_message = "Validation syntaxique des annotations et placeholders"
    await db.commit()

    # 2. Résoudre le nom du dossier pour les chemins
    dossier_name = await _get_dossier_name(dossier_id, db)

    # 3. Vérifier que le PEA existe
    in_dir = _step_in(dossier_name, 4)
    pea_path = os.path.join(in_dir, "pea.docx")
    if not os.path.isfile(pea_path):
        await workflow_engine.fail_step(dossier_id, 4, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="PEA/PAA non trouvé — importez d'abord le fichier dans la section Fichiers d'entrée",
        )

    # 4. Créer le répertoire de sortie
    out_dir = _step_out(dossier_name, 4)
    os.makedirs(out_dir, exist_ok=True)
    pre_path = os.path.join(out_dir, "pre.docx")

    # 5. Copier le PEA → pre.docx (document de travail)
    # Le PEA est maintenant le TRE complet annoté par l'expert (plus besoin de recoller l'en-tête)
    shutil.copy2(pea_path, pre_path)

    # 5. Charger les placeholders de réquisition (Step 1)
    placeholders_csv_path = os.path.join(
        _step_out(dossier_name, 1), "placeholders.csv"
    )
    placeholders: dict[str, str] = {}
    if os.path.isfile(placeholders_csv_path):
        with open(placeholders_csv_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("nom_placeholder"):
                    continue
                parts = line.split(";", 1)
                if len(parts) == 2 and parts[1].strip():
                    placeholders[parts[0].strip()] = parts[1].strip()
        logger.info(
            "[Step4] Dossier %d — %d placeholders chargés",
            dossier_id, len(placeholders),
        )

    # 5b. Générer les alias pour couvrir les anciennes conventions de nommage
    # Le TRE peut encore utiliser les anciens noms (mec, tribunal, etc.)
    _ALIAS_MAP = {
        # Nouveaux noms → anciens noms (pour TRE pas encore mis à jour)
        "nom_pex": ["nom_mec", "mec_nom"],
        "prenom_pex": ["prenom_mec", "mec_prenom"],
        "genre_pex": ["genre_mec", "mec_genre"],
        "date_naissance_pex": ["date_naissance_mec"],
        "ville_naissance_pex": ["ville_naissance_mec"],
        "nom_expert": ["expert_nom"],
        "prenom_expert": ["expert_prenom"],
        "titre_expert": ["expert_titre"],
        "genre_expert": ["expert_genre"],
        "titre_expertise": ["nom_expertise", "expertise_nom"],
        "objet_mission": ["mission_objet"],
        "nom_tribunal": ["tribunal", "tribunal_nom"],
        "ville_tribunal": ["ville_juridiction", "requerant_ville"],
        "nom_magistrat": ["magistrat"],
        "requerant_nom": ["nom_requerant"],
        "requerant_prenom": ["prenom_requerant"],
        "requerant_titre": ["titre_requerant"],
    }
    # Pour chaque placeholder connu, générer les alias
    aliases_to_add: dict[str, str] = {}
    for canonical, alias_list in _ALIAS_MAP.items():
        if canonical in placeholders:
            for alias in alias_list:
                if alias not in placeholders:
                    aliases_to_add[alias] = placeholders[canonical]
        else:
            # Chercher si un alias existe et créer le canonical
            for alias in alias_list:
                if alias in placeholders:
                    aliases_to_add[canonical] = placeholders[alias]
                    break
    placeholders.update(aliases_to_add)
    if aliases_to_add:
        logger.info(
            "[Step4] Dossier %d — %d alias de placeholders ajoutés",
            dossier_id, len(aliases_to_add),
        )

    # 6. Parser le PEA pour extraire les annotations
    parser = TREParser()
    try:
        pea_parse_result = parser.parse(pre_path)
    except Exception as exc:
        logger.error("[Step4] Erreur parsing PEA : %s", exc)
        await workflow_engine.fail_step(dossier_id, 4, db)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Erreur lors du parsing du PEA : {exc}",
        )

    # Classer les annotations par type — utiliser un index séquentiel pour gérer les doublons
    annotations_dires: list[tuple[str, str]] = []  # (full_key, content)
    annotations_analyse: list[tuple[str, str]] = []
    annotations_verbatim: list[tuple[str, str]] = []

    for annot in pea_parse_result.annotations:
        if annot.type == "debut_tpe":
            continue
        full_key = f"{annot.type}_{annot.suffix}" if annot.suffix else annot.type
        if annot.type == "dires":
            annotations_dires.append((full_key, annot.content))
        elif annot.type == "analyse":
            annotations_analyse.append((full_key, annot.content))
        elif annot.type == "verbatim":
            annotations_verbatim.append((full_key, annot.content))

    logger.info(
        "[Step4] Dossier %d — Annotations : %d dires, %d analyse, %d verbatim",
        dossier_id, len(annotations_dires), len(annotations_analyse),
        len(annotations_verbatim),
    )

    # ──────────────────────────────────────────────────────────────────────
    # ÉTAPE 2/5 : Reformulation LLM des @dires et @analyse (batch)
    # ──────────────────────────────────────────────────────────────────────
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 2
    step.progress_message = f"Reformulation @dires (batch {len(annotations_dires)}) ⏳ LLM"
    await db.commit()

    llm = LLMService()

    # Reformuler tous les @dires en un seul appel LLM
    reformulated_dires: list[str] = []
    if annotations_dires:
        non_empty_dires = [(k, c) for k, c in annotations_dires if c.strip()]
        if non_empty_dires:
            try:
                reformulated_dires = await llm.reformuler_dires_batch(non_empty_dires)
            except Exception as exc:
                logger.warning("[Step4] Erreur batch reformulation dires : %s — fallback séquentiel", exc)
                # Fallback séquentiel si le batch échoue
                reformulated_dires = []
                for key, content in non_empty_dires:
                    try:
                        reformulated_dires.append(await llm.reformuler_dires(content))
                    except Exception:
                        reformulated_dires.append(content)
        # Réinsérer les vides aux bonnes positions
        full_dires: list[str] = []
        non_empty_idx = 0
        for key, content in annotations_dires:
            if content.strip():
                full_dires.append(reformulated_dires[non_empty_idx] if non_empty_idx < len(reformulated_dires) else content)
                non_empty_idx += 1
            else:
                full_dires.append("")
        reformulated_dires = full_dires

    # Reformuler tous les @analyse en un seul appel LLM
    step = await _get_step(dossier_id, 4, db)
    step.progress_message = f"Reformulation @analyse (batch {len(annotations_analyse)}) ⏳ LLM"
    await db.commit()

    reformulated_analyse: list[str] = []
    if annotations_analyse:
        non_empty_analyse = [(k, c) for k, c in annotations_analyse if c.strip()]
        if non_empty_analyse:
            try:
                reformulated_analyse = await llm.reformuler_analyse_batch(non_empty_analyse)
            except Exception as exc:
                logger.warning("[Step4] Erreur batch reformulation analyse : %s — fallback séquentiel", exc)
                reformulated_analyse = []
                for key, content in non_empty_analyse:
                    try:
                        reformulated_analyse.append(await llm.reformuler_analyse(content))
                    except Exception:
                        reformulated_analyse.append(content)
        # Réinsérer les vides aux bonnes positions
        full_analyse: list[str] = []
        non_empty_idx = 0
        for key, content in annotations_analyse:
            if content.strip():
                full_analyse.append(reformulated_analyse[non_empty_idx] if non_empty_idx < len(reformulated_analyse) else content)
                non_empty_idx += 1
            else:
                full_analyse.append("")
        reformulated_analyse = full_analyse

    # Construire un dictionnaire clé → texte reformulé (pour @resume et @cite)
    reformulated_by_key: dict[str, str] = {}
    for i, (key, _) in enumerate(annotations_dires):
        if i < len(reformulated_dires) and reformulated_dires[i]:
            reformulated_by_key[key] = reformulated_dires[i]
    for i, (key, _) in enumerate(annotations_analyse):
        if i < len(reformulated_analyse) and reformulated_analyse[i]:
            reformulated_by_key[key] = reformulated_analyse[i]

    # ──────────────────────────────────────────────────────────────────────
    # ÉTAPE 3/5 : Résolution des @resume ⏳ LLM
    # ──────────────────────────────────────────────────────────────────────
    # @resume concatène les sections référencées et génère un résumé LLM
    # Format : @resume dires_section_7.1 dires_section_7.2 ...@
    # On parse les clés, on concatène les textes reformulés, on résume via LLM

    # Collecter les @resume depuis le parsing
    annotations_resume: list[tuple[str, str]] = []
    for annot in pea_parse_result.annotations:
        if annot.type == "resume":
            full_key = f"resume_{annot.suffix}" if annot.suffix else "resume"
            annotations_resume.append((full_key, annot.content))

    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 3
    step.progress_message = f"Résolution des @resume {len(annotations_resume)} ⏳ LLM"
    await db.commit()

    # Résoudre chaque @resume
    resolved_resumes: list[str] = []
    for i, (key, content) in enumerate(annotations_resume):
        # Parser les clés référencées (séparées par espaces ou virgules)
        ref_keys = [k.strip().strip(",") for k in content.split() if k.strip().strip(",")]
        # Concaténer les textes reformulés des sections citées
        concat_parts = []
        for ref_key in ref_keys:
            if ref_key in reformulated_by_key:
                concat_parts.append(reformulated_by_key[ref_key])
            else:
                concat_parts.append(f"[Section {ref_key} non trouvée]")

        if concat_parts:
            concat_text = "\n\n".join(concat_parts)
            # Résumer via LLM
            try:
                summary = await llm.chat(
                    [{"role": "user", "content": concat_text}],
                    system_prompt=(
                        "Tu es un assistant spécialisé en rédaction de rapports d'expertise judiciaire.\n"
                        "Résume le texte suivant de manière concise et professionnelle, "
                        "en conservant les éléments cliniques essentiels.\n"
                        "Réponds uniquement avec le résumé, sans commentaire."
                    ),
                )
                resolved_resumes.append(summary)
            except Exception as exc:
                logger.warning("[Step4] Erreur résumé @resume %s : %s", key, exc)
                resolved_resumes.append(concat_text)  # fallback : concaténation sans résumé
        else:
            resolved_resumes.append("[Aucune section référencée trouvée]")

        # Mise à jour progression
        step = await _get_step(dossier_id, 4, db)
        step.progress_message = f"Résolution @resume {i + 1}/{len(annotations_resume)} ⏳ LLM"
        await db.commit()

    # ──────────────────────────────────────────────────────────────────────
    # ÉTAPE 4/5 : Résolution @question, @reference, @cite
    # ──────────────────────────────────────────────────────────────────────
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 4
    step.progress_message = "Résolution @question, @reference, @cite"
    await db.commit()

    # ──────────────────────────────────────────────────────────────────────
    # ÉTAPE 5/5 : Reconstitution du document et substitution in-place
    # ──────────────────────────────────────────────────────────────────────
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 5
    step.progress_message = "Reconstitution et substitution dans le document Word"
    await db.commit()

    # --- Stratégie ---
    # Le document pre.docx est maintenant reconstitué : en-tête TRE + contenu PEA.
    # On parcourt les paragraphes pour :
    # 1. Identifier les blocs d'annotations multi-paragraphes (début @type ... fin @)
    # 2. Remplacer chaque bloc par le texte reformulé
    # 3. Substituer les <<placeholders>> dans tout le document

    doc = DocxDocument(pre_path)
    paragraphs = doc.paragraphs

    # Regex patterns
    placeholder_re = re_module.compile(r"<<(\w+)>>")
    debut_tpe_re = re_module.compile(r"^\s*@debut_tpe@\s*$")
    # Annotation complète sur une ligne : @type contenu@ (greedy sur le type, lazy sur le contenu)
    # Le type peut contenir des underscores et des points (ex: dires_3.1.2)
    annotation_full_re = re_module.compile(
        r"^\s*@(\w+(?:_[\w.]+)*)\s+(.*?)\s*@\s*$", re_module.DOTALL
    )
    # Ouverture d'annotation au début d'un paragraphe : @type_suffix suivi de texte (pas de @ fermant)
    annotation_open_re = re_module.compile(r"^\s*@(\w+(?:_[\w.]+)*)\s+(.*)")
    # Fermeture : le paragraphe se termine par un @ isolé (pas précédé d'un autre @)
    annotation_close_suffix = re_module.compile(r"^(.*?)\s*@\s*$", re_module.DOTALL)

    # Compteurs séquentiels pour les annotations reformulées
    dires_idx = 0
    analyse_idx = 0
    resume_idx = 0

    def _get_replacement(annot_type: str, content: str) -> str:
        """Retourne le texte de remplacement pour une annotation."""
        nonlocal dires_idx, analyse_idx, resume_idx

        if annot_type.startswith("dires"):
            if dires_idx < len(reformulated_dires) and reformulated_dires[dires_idx]:
                result = f"Dires :\n{reformulated_dires[dires_idx]}"
            else:
                result = f"Dires :\n{content}" if content.strip() else ""
            dires_idx += 1
            return result
        elif annot_type.startswith("analyse"):
            if analyse_idx < len(reformulated_analyse) and reformulated_analyse[analyse_idx]:
                result = f"Analyse :\n{reformulated_analyse[analyse_idx]}"
            else:
                result = f"Analyse :\n{content}" if content.strip() else ""
            analyse_idx += 1
            return result
        elif annot_type.startswith("verbatim"):
            return f"\u00ab {content} \u00bb" if content.strip() else ""
        elif annot_type.startswith("remplir"):
            # @remplir description : texte@ → garder seulement après ":"
            if ":" in content:
                return content.split(":", 1)[1].strip()
            return content.strip()
        elif annot_type == "question":
            # @question N@ → substituer depuis placeholders
            q_num = content.strip()
            q_key = f"question_{q_num}"
            return placeholders.get(q_key, f"[Question {q_num} non trouvée]")
        elif annot_type.startswith("resume"):
            # @resume refs@ → texte résumé (résolu à l'étape 3)
            if resume_idx < len(resolved_resumes):
                result = resolved_resumes[resume_idx]
                resume_idx += 1
                return result
            return "[Résumé non résolu]"
        elif annot_type.startswith("reference"):
            return f"(cf. {content})"
        elif annot_type.startswith("cite"):
            # @cite dires_section_6.1@ → insérer le contenu reformulé de la section
            ref_key = content.strip()
            if ref_key in reformulated_by_key:
                return reformulated_by_key[ref_key]
            return f"[Citation {ref_key} non trouvée]"
        elif annot_type == "debut_tpe":
            return ""
        else:
            # Annotation personnalisée ou inconnue — garder le contenu
            return content.strip() if content.strip() else ""

    def _replace_paragraph_text(paragraph, new_text: str) -> None:
        """Remplace le texte d'un paragraphe en préservant le style du paragraphe.

        Stratégie : concaténer tout le texte existant des runs, puis remplacer.
        On garde le premier run avec le nouveau texte et on vide les autres.
        Cela préserve le style du paragraphe (Heading, Normal, etc.) et le
        formatage de base du premier run.
        """
        runs = paragraph.runs
        if not runs:
            # Pas de runs — créer le texte directement
            if new_text:
                paragraph.add_run(new_text)
            return
        # Mettre le nouveau texte dans le premier run, vider les autres
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""

    def _clear_paragraph(paragraph) -> None:
        """Vide un paragraphe."""
        _replace_paragraph_text(paragraph, "")

    # --- Passe 1 : traiter les annotations (single-line et multi-paragraphes) ---
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i].text
        if not text or not text.strip():
            i += 1
            continue

        # 1. Supprimer @debut_tpe@
        if debut_tpe_re.match(text):
            _clear_paragraph(paragraphs[i])
            i += 1
            continue

        # 2. Annotation complète sur une seule ligne : @type_suffix contenu@
        full_match = annotation_full_re.match(text)
        if full_match:
            annot_type = full_match.group(1)
            content = full_match.group(2).strip()
            replacement = _get_replacement(annot_type, content)
            _replace_paragraph_text(paragraphs[i], replacement)
            i += 1
            continue

        # 3. Ouverture d'annotation multi-paragraphe : @type_suffix contenu...
        open_match = annotation_open_re.match(text)
        if open_match:
            annot_type = open_match.group(1)
            first_content = open_match.group(2)

            # Vérifier si la fermeture est dans ce même paragraphe (cas rare avec texte avant)
            # Non — on cherche la fermeture dans les paragraphes suivants
            content_parts = [first_content]
            close_found = False
            j = i + 1
            while j < len(paragraphs):
                p_text = paragraphs[j].text
                close_match = annotation_close_suffix.match(p_text)
                if close_match:
                    content_parts.append(close_match.group(1))
                    close_found = True
                    break
                else:
                    content_parts.append(p_text)
                j += 1

            if close_found:
                full_content = "\n".join(content_parts).strip()
                replacement = _get_replacement(annot_type, full_content)
                _replace_paragraph_text(paragraphs[i], replacement)
                # Vider les paragraphes intermédiaires et de fermeture
                for k in range(i + 1, j + 1):
                    _clear_paragraph(paragraphs[k])
                i = j + 1
                continue
            else:
                # Pas de fermeture trouvée — laisser tel quel (annotation mal formée)
                logger.warning(
                    "[Step4] Annotation @%s non fermée au paragraphe %d",
                    annot_type, i,
                )
                i += 1
                continue

        # 4. Pas une annotation — passer au suivant
        i += 1

    # --- Passe 2 : substituer les <<placeholders>> dans tous les paragraphes ---
    def _substitute_placeholders_in_para(paragraph) -> None:
        """Remplace les <<placeholder>> par leurs valeurs dans un paragraphe."""
        text = paragraph.text
        if not text or "<<" not in text:
            return
        new_text = placeholder_re.sub(
            lambda m: placeholders.get(m.group(1), f"[{m.group(1)}]"),
            text,
        )
        if new_text != text:
            _replace_paragraph_text(paragraph, new_text)

    for paragraph in paragraphs:
        _substitute_placeholders_in_para(paragraph)

    # Aussi traiter les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _substitute_placeholders_in_para(paragraph)

    # Aussi traiter les en-têtes et pieds de page
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    _substitute_placeholders_in_para(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    _substitute_placeholders_in_para(paragraph)

    # Sauvegarder le .docx modifié
    doc.save(pre_path)

    # ──────────────────────────────────────────────────────────────────────
    # Sauvegarde en base
    # ──────────────────────────────────────────────────────────────────────

    step = await _get_step(dossier_id, 4, db)

    # Supprimer les anciens fichiers PRE (ré-exécution) mais garder le PEA et le DAC
    for old_file in list(step.files):
        if old_file.file_type == "re_projet":
            await db.delete(old_file)
    await db.flush()

    pre_size = os.path.getsize(pre_path)
    db.add(StepFile(
        step_id=step.id,
        filename="pre.docx",
        file_path=pre_path,
        file_type="re_projet",
        file_size=pre_size,
    ))

    # Marquer step4 comme "fait"
    await workflow_engine.execute_step(dossier_id, 4, db)
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = None
    step.progress_total = None
    step.progress_message = None
    await db.commit()

    logger.info("[Step4] Dossier %d — PRE généré avec succès (%d octets)", dossier_id, pre_size)

    return Step2UploadResponse(
        message="PRE généré avec succès",
        filenames=["pre.docx"],
    )


# ---- Step4 : POST /api/dossiers/{id}/step4/generate-dac -----------------

@router.post(
    "/{dossier_id}/step4/generate-dac",
    response_model=Step2UploadResponse,
)
async def step4_generate_dac(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Génération optionnelle du DAC (Document d'Analyse Contradictoire).

    Peut être lancé après la génération du PRE. Utilise le PRE existant
    et le PEA pour produire une analyse contradictoire via LLM + RAG.
    """

    # 1. Vérifier l'accès au step4
    await workflow_engine.require_step_access(dossier_id, 4, db)

    # 2. Résoudre le nom du dossier
    dossier_name = await _get_dossier_name(dossier_id, db)
    out_dir = _step_out(dossier_name, 4)
    in_dir = _step_in(dossier_name, 4)

    # 3. Vérifier que le PRE existe
    pre_path = os.path.join(out_dir, "pre.docx")
    if not os.path.isfile(pre_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le PRE n'a pas encore été généré — lancez d'abord l'opération principale",
        )

    # 4. Lire le PRE (texte brut pour le DAC)
    from docx import Document as DocxDocument
    pre_doc = DocxDocument(pre_path)
    pre_content = "\n".join(p.text for p in pre_doc.paragraphs if p.text.strip())

    # 5. Lire le PEA (texte brut)
    pea_path = os.path.join(in_dir, "pea.docx")
    pea_text = ""
    if os.path.isfile(pea_path):
        with open(pea_path, "rb") as f:
            pea_text = f.read().decode("utf-8", errors="replace")

    # 6. Mettre à jour la progression
    step = await _get_step(dossier_id, 4, db)
    step.progress_current = 1
    step.progress_total = 1
    step.progress_message = "Génération du DAC ⏳ LLM"
    await db.commit()

    # 7. Récupérer le domaine
    result = await db.execute(select(LocalConfig).limit(1))
    config = result.scalar_one_or_none()
    domaine = config.domaine if config else "psychologie"

    # 8. Récupérer le contexte RAG
    rag = RAGService()
    rag_context = ""
    try:
        rag_docs = await rag.search(
            query="contestation expertise méthodologie biais déontologie jurisprudence",
            collection=f"corpus_{domaine}",
            limit=10,
        )
        if rag_docs:
            rag_context = "\n\n".join(doc.content for doc in rag_docs)
    except Exception as exc:
        logger.warning(
            "[Step4/DAC] Dossier %d — Contexte RAG indisponible : %s",
            dossier_id, exc,
        )

    # 9. Générer le DAC via LLM
    llm = LLMService()
    try:
        dac_content = await llm.generer_dac(pea_text, pre_content, rag_context)
    except Exception as exc:
        logger.error("Erreur LLM lors de la génération du DAC : %s", exc)
        step = await _get_step(dossier_id, 4, db)
        step.progress_current = None
        step.progress_total = None
        step.progress_message = None
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Service LLM indisponible — essayez de redémarrer le conteneur judi-llm",
        )

    # 10. Sauvegarder le DAC
    os.makedirs(out_dir, exist_ok=True)
    dac_path = os.path.join(out_dir, "dac.docx")
    _save_markdown_as_docx(dac_content, dac_path, "Document d'Analyse Contradictoire")

    # Supprimer l'ancien DAC en base s'il existe
    step = await _get_step(dossier_id, 4, db)
    for old_file in list(step.files):
        if old_file.file_type == "re_projet_auxiliaire":
            await db.delete(old_file)
    await db.flush()

    dac_size = os.path.getsize(dac_path)
    db.add(StepFile(
        step_id=step.id,
        filename="dac.docx",
        file_path=dac_path,
        file_type="re_projet_auxiliaire",
        file_size=dac_size,
    ))

    step.progress_current = None
    step.progress_total = None
    step.progress_message = None
    await db.commit()

    logger.info("[Step4/DAC] Dossier %d — DAC généré avec succès", dossier_id)

    return Step2UploadResponse(
        message="DAC généré avec succès",
        filenames=["dac.docx"],
    )


# ---- Step4 : POST /api/dossiers/{id}/step4/validate ---------------------

@router.post(
    "/{dossier_id}/step4/validate",
    response_model=Step2ValidateResponse,
)
async def step4_validate(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Valide le Step 4 (verrouillage) et autorise l'accès au Step 5.

    Valide : Exigence 8.4
    """

    await workflow_engine.validate_step(dossier_id, 4, db)
    await db.commit()

    return Step2ValidateResponse(message="Step 4 validé avec succès")


# ---- Step5 : POST /api/dossiers/{id}/step5/execute -----------------------

@router.post(
    "/{dossier_id}/step5/execute",
    response_model=Step3ExecuteResponse,
    status_code=status.HTTP_201_CREATED,
)
async def step5_execute(
    dossier_id: int,
    file: UploadFile = None,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Finalisation, révision linguistique et archivage du dossier.

    Nouveau flux TRE-centré :
    1. Upload du REF (rapport final ajusté par l'expert)
    2. Service de révision linguistique (préservation verbatim)
    3. Génération du ZIP (tous les fichiers sauf archive/)
    4. Génération du timbre.txt (métadonnées + hash SHA-256)
    5. Placement dans <nom-dossier>/archive/

    Valide : Requirement 5 (Finalisation et archivage)
    """
    import hashlib
    import io
    import zipfile
    from datetime import datetime, UTC

    from services.revision_service import RevisionService
    from services.file_paths import create_archive_dir, dossier_root as get_dossier_root

    # 1. Vérifier l'accès au step5
    await workflow_engine.require_step_access(dossier_id, 5, db)

    # 1b. Marquer le step comme "en_cours"
    await workflow_engine.start_step(dossier_id, 5, db)
    await db.commit()

    # Résoudre le nom du dossier pour les chemins
    dossier_name = await _get_dossier_name(dossier_id, db)

    in_dir = _step_in(dossier_name, 5)
    out_dir = _step_out(dossier_name, 5)
    os.makedirs(in_dir, exist_ok=True)
    os.makedirs(out_dir, exist_ok=True)

    # 2. Sauvegarder le REF uploadé (rapport final ajusté)
    ref_text = ""
    if file and file.filename:
        ref_bytes = await file.read()
        ref_path = os.path.join(in_dir, "ref.docx")
        with open(ref_path, "wb") as f:
            f.write(ref_bytes)

        # Extraire le texte du REF pour la révision
        try:
            from docx import Document as DocxDocument
            ref_doc = DocxDocument(ref_path)
            ref_text = "\n".join(p.text for p in ref_doc.paragraphs if p.text.strip())
        except Exception as exc:
            logger.warning("[Step5] Impossible d'extraire le texte du REF : %s", exc)

    # 3. Service de révision linguistique (optionnel, si texte disponible)
    revision_corrections = []
    if ref_text.strip():
        logger.info("[Step5] Dossier %d — Révision linguistique en cours…", dossier_id)
        revision_service = RevisionService()
        try:
            revision_result = await revision_service.revise(ref_text)
            revision_corrections = [
                {"original": c.original, "corrected": c.corrected, "position": c.position}
                for c in revision_result.corrections
            ]
            logger.info(
                "[Step5] Dossier %d — %d corrections, %d verbatim préservés",
                dossier_id, len(revision_corrections), revision_result.verbatim_count,
            )
        except Exception as exc:
            logger.warning("[Step5] Révision échouée : %s — on continue sans", exc)

    # 4. Construire le ZIP avec TOUS les fichiers du dossier (sauf archive/)
    root_dir = get_dossier_root(dossier_name)
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for dirpath, dirnames, filenames in os.walk(root_dir):
            # Exclure le répertoire archive/
            if "archive" in dirnames:
                dirnames.remove("archive")
            for fname in filenames:
                full_path = os.path.join(dirpath, fname)
                arc_name = os.path.relpath(full_path, root_dir)
                zf.write(full_path, arc_name)

    zip_bytes = zip_buffer.getvalue()

    # 5. Générer le hash SHA-256 du ZIP
    sha256_hash = hashlib.sha256(zip_bytes).hexdigest()

    # 6. Créer le répertoire archive et sauvegarder le ZIP
    archive_path = create_archive_dir(dossier_name)

    zip_filename = f"{dossier_name}.zip"
    zip_path = os.path.join(archive_path, zip_filename)
    with open(zip_path, "wb") as f:
        f.write(zip_bytes)

    # 7. Générer le fichier timbre.txt avec les métadonnées
    # Charger les placeholders pour les métadonnées
    placeholders_csv_path = os.path.join(_step_out(dossier_name, 1), "placeholders.csv")
    meta: dict[str, str] = {}
    if os.path.isfile(placeholders_csv_path):
        with open(placeholders_csv_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("nom_placeholder"):
                    continue
                parts = line.split(";", 1)
                if len(parts) == 2 and parts[1].strip():
                    meta[parts[0].strip()] = parts[1].strip()

    now = datetime.now(UTC)
    timbre_filename = f"{dossier_name}-timbre.txt"
    timbre_path = os.path.join(archive_path, timbre_filename)
    with open(timbre_path, "w", encoding="utf-8") as f:
        f.write(f"=== TIMBRE D'ARCHIVAGE — {dossier_name} ===\n\n")
        f.write(f"Date d'archivage : {now.strftime('%d/%m/%Y %H:%M:%S UTC')}\n\n")
        f.write("--- Contexte expertise ---\n")
        f.write(f"Demandeur nom : {meta.get('nom_demandeur', '')}\n")
        f.write(f"Demandeur prénom : {meta.get('prenom_demandeur', '')}\n")
        f.write(f"Demandeur adresse : {meta.get('adresse_demandeur', '')}\n")
        f.write(f"Demande date : {meta.get('date_ordonnance', '')}\n")
        f.write(f"Tribunal nom : {meta.get('tribunal', '')}\n")
        f.write(f"Tribunal adresse : {meta.get('ville_juridiction', '')}\n")
        f.write(f"Demande référence : {meta.get('reference_dossier', '')}\n")
        f.write(f"MEC nom : {meta.get('nom_mec', meta.get('nom_defendeur', ''))}\n")
        f.write(f"MEC prénom : {meta.get('prenom_mec', meta.get('prenom_defendeur', ''))}\n")
        f.write(f"MEC adresse : {meta.get('adresse_mec', '')}\n")
        f.write(f"Expert nom : {meta.get('nom_expert', '')}\n")
        f.write(f"Expert prénom : {meta.get('prenom_expert', '')}\n")
        f.write(f"Expert adresse : {meta.get('adresse_expert', '')}\n")
        f.write(f"\n--- Archive ---\n")
        f.write(f"Fichier archive : {zip_filename}\n")
        f.write(f"SHA-256 : {sha256_hash}\n")
        f.write(f"Taille : {len(zip_bytes)} octets\n")
        if revision_corrections:
            f.write(f"\n--- Révision linguistique ---\n")
            f.write(f"Corrections appliquées : {len(revision_corrections)}\n")

    logger.info(
        "[Step5] Dossier %d — Archive %s (%d octets), hash %s",
        dossier_id, zip_filename, len(zip_bytes), sha256_hash,
    )

    # 8. Créer les entrées StepFile en base
    step = await _get_step(dossier_id, 5, db)

    # Supprimer les anciens StepFile
    for old_file in list(step.files):
        await db.delete(old_file)
    await db.flush()

    if file and file.filename:
        db.add(StepFile(
            step_id=step.id,
            filename="ref.docx",
            file_path=os.path.join(in_dir, "ref.docx"),
            file_type="rapport_final",
            file_size=len(ref_bytes),
        ))

    db.add(StepFile(
        step_id=step.id,
        filename=zip_filename,
        file_path=zip_path,
        file_type="archive_zip",
        file_size=len(zip_bytes),
    ))
    db.add(StepFile(
        step_id=step.id,
        filename=timbre_filename,
        file_path=timbre_path,
        file_type="timbre",
        file_size=os.path.getsize(timbre_path),
    ))

    # 9. Marquer step5 comme "fait"
    await workflow_engine.execute_step(dossier_id, 5, db)
    await db.commit()

    return Step3ExecuteResponse(
        message=f"Archive générée — Hash SHA-256 : {sha256_hash}",
        filenames=[zip_filename, timbre_filename],
    )


# ---- Step5 : GET /api/dossiers/{id}/step5/download/{doc_type} ------------

@router.get(
    "/{dossier_id}/step5/download/{doc_type}",
)
async def step5_download(
    dossier_id: int,
    doc_type: str,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Télécharge le REF-Projet généré pour step3.

    Valide : Exigence 4.4
    """

    # Vérifier l'accès à l'étape
    await workflow_engine.require_step_access(dossier_id, 5, db)

    if doc_type != "ref_projet":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Type de document invalide — utilisez 'ref_projet'",
        )

    dossier_name = await _get_dossier_name(dossier_id, db)
    file_path = os.path.join(_step_out(dossier_name, 5), "ref_projet.docx")
    if not os.path.isfile(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Fichier REF-Projet non trouvé — lancez d'abord l'exécution du Step3",
        )

    return FileResponse(
        path=file_path,
        filename="ref_projet.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---- Step5 : POST /api/dossiers/{id}/step5/validate ---------------------

@router.post(
    "/{dossier_id}/step5/validate",
    response_model=Step3ValidateResponse,
)
async def step5_validate(
    dossier_id: int,
    _user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Valide le Step 5 (verrouillage).

    Valide : Exigences 9.5, 9.6
    """

    await workflow_engine.validate_step(dossier_id, 5, db)
    await db.commit()

    return Step3ValidateResponse(message="Step 5 validé avec succès")
