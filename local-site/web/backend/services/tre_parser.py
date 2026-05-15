"""
Judi-Expert — Parser de documents TRE (Template de Rapport d'Expertise).

Parse les fichiers .docx TRE pour extraire les méta-instructions :
- Placeholders : <<nom>> (variables à remplacer)
- Annotations : @type contenu@ (instructions structurelles)

Gère les annotations multi-paragraphes et les annotations personnalisées.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------


@dataclass
class Placeholder:
    """Représente un placeholder <<nom>> dans le TRE.

    Attributes:
        name: Nom du placeholder (ex: "nom_expert", "question_1").
        position: Index du paragraphe dans le document .docx.
    """

    name: str
    position: int


@dataclass
class Annotation:
    """Représente une annotation @type contenu@ dans le TRE.

    Attributes:
        type: Type d'annotation ("dires", "analyse", "verbatim", etc.).
        suffix: Suffixe optionnel (ex: "fratrie", "2.1.3", "").
        content: Contenu textuel entre les marqueurs @...@.
        position: Index du paragraphe de début dans le document .docx.
        is_custom: True si l'annotation est personnalisée (@/mon_annotation).
    """

    type: str
    suffix: str
    content: str
    position: int
    is_custom: bool


@dataclass
class TREParseResult:
    """Résultat du parsing d'un document TRE.

    Attributes:
        placeholders: Liste des placeholders trouvés.
        annotations: Liste des annotations trouvées.
        debut_tpe_position: Position du marqueur @debut_tpe@, ou None.
        errors: Liste des erreurs rencontrées lors du parsing.
    """

    placeholders: list[Placeholder] = field(default_factory=list)
    annotations: list[Annotation] = field(default_factory=list)
    debut_tpe_position: int | None = None
    errors: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# TREParser
# ---------------------------------------------------------------------------


class TREParser:
    """Parser de documents TRE (.docx).

    Extrait les placeholders et annotations d'un Template de Rapport
    d'Expertise au format Word (.docx).
    """

    # Types d'annotations prédéfinis
    PREDEFINED_TYPES = frozenset(
        {"dires", "analyse", "verbatim", "question", "reference", "cite", "debut_tpe", "remplir", "resume"}
    )

    # Patterns regex
    PLACEHOLDER_RE = re.compile(r"<<(\w+)>>")
    DEBUT_TPE_RE = re.compile(r"^\s*@debut_tpe@\s*$")
    # Ouverture d'annotation : @type ou @type_suffix ou @/custom
    ANNOTATION_OPEN_RE = re.compile(r"@(/?\w+(?:_[\w.]+)*)\s")
    # Fermeture d'annotation : contenu se terminant par @
    ANNOTATION_CLOSE_RE = re.compile(r"(.*?)\s*@\s*$", re.DOTALL)
    # Annotation sur une seule ligne : @type contenu@
    ANNOTATION_SINGLE_RE = re.compile(r"@(/?\w+(?:_[\w.]+)*)\s+(.*?)\s*@", re.DOTALL)

    # Validation snake_case pour les noms de placeholders
    SNAKE_CASE_RE = re.compile(r"^[a-z][a-z0-9]*(_[a-z0-9]+)*$")

    def parse(self, docx_path: str) -> TREParseResult:
        """Parse le TRE et extrait toutes les méta-instructions.

        Args:
            docx_path: Chemin vers le fichier .docx du TRE.

        Returns:
            TREParseResult contenant placeholders, annotations et erreurs.
        """
        result = TREParseResult()
        doc = DocxDocument(docx_path)

        # État pour les annotations multi-paragraphes
        open_annotation: dict | None = None

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text

            # 1. Vérifier le marqueur @debut_tpe@
            if self.DEBUT_TPE_RE.match(text):
                if result.debut_tpe_position is not None:
                    result.errors.append(
                        f"Paragraphe {idx} : marqueur @debut_tpe@ en double "
                        f"(déjà trouvé au paragraphe {result.debut_tpe_position})."
                    )
                else:
                    result.debut_tpe_position = idx
                    result.annotations.append(
                        Annotation(
                            type="debut_tpe",
                            suffix="",
                            content="",
                            position=idx,
                            is_custom=False,
                        )
                    )
                continue

            # 2. Si une annotation multi-paragraphe est ouverte
            if open_annotation is not None:
                close_match = self.ANNOTATION_CLOSE_RE.match(text)
                if close_match:
                    # Fermeture de l'annotation
                    open_annotation["content_parts"].append(close_match.group(1))
                    full_content = "\n".join(open_annotation["content_parts"])
                    annotation_type, suffix, is_custom = self._parse_annotation_key(
                        open_annotation["key"]
                    )
                    result.annotations.append(
                        Annotation(
                            type=annotation_type,
                            suffix=suffix,
                            content=full_content.strip(),
                            position=open_annotation["position"],
                            is_custom=is_custom,
                        )
                    )
                    open_annotation = None
                else:
                    # Continuation de l'annotation multi-paragraphe
                    open_annotation["content_parts"].append(text)
                continue

            # 3. Vérifier les placeholders <<nom>>
            placeholder_matches = self.PLACEHOLDER_RE.findall(text)
            for name in placeholder_matches:
                result.placeholders.append(
                    Placeholder(name=name, position=idx)
                )

            # 4. Vérifier les annotations (single-line d'abord)
            single_match = self.ANNOTATION_SINGLE_RE.search(text)
            if single_match:
                key = single_match.group(1)
                content = single_match.group(2)
                annotation_type, suffix, is_custom = self._parse_annotation_key(key)
                result.annotations.append(
                    Annotation(
                        type=annotation_type,
                        suffix=suffix,
                        content=content.strip(),
                        position=idx,
                        is_custom=is_custom,
                    )
                )
                continue

            # 5. Vérifier ouverture d'annotation multi-paragraphe
            open_match = self.ANNOTATION_OPEN_RE.search(text)
            if open_match:
                key = open_match.group(1)
                # Le contenu commence après le marqueur d'ouverture
                content_start = text[open_match.end():]
                open_annotation = {
                    "key": key,
                    "position": idx,
                    "content_parts": [content_start],
                }

        # Vérifier si une annotation est restée ouverte
        if open_annotation is not None:
            result.errors.append(
                f"Paragraphe {open_annotation['position']} : annotation "
                f"@{open_annotation['key']} non fermée (@ manquant)."
            )

        return result

    def validate(self, result: TREParseResult) -> list[str]:
        """Valide la structure du TRE parsé.

        Args:
            result: Résultat du parsing à valider.

        Returns:
            Liste des erreurs de validation.
        """
        errors: list[str] = []

        # Vérifier la présence du marqueur @debut_tpe@
        if result.debut_tpe_position is None:
            errors.append("Le marqueur @debut_tpe@ est absent du document.")

        # Vérifier les annotations non fermées (déjà dans result.errors)
        for error in result.errors:
            if "non fermée" in error:
                errors.append(error)

        # Vérifier que les noms de placeholders sont en snake_case
        for placeholder in result.placeholders:
            if not self.SNAKE_CASE_RE.match(placeholder.name):
                errors.append(
                    f"Paragraphe {placeholder.position} : le placeholder "
                    f"<<{placeholder.name}>> n'est pas en snake_case."
                )

        return errors

    def extract_pe(self, docx_path: str, questions: dict[str, str]) -> bytes:
        """Extrait le PE (Plan d'Entretien) depuis @debut_tpe@ jusqu'à la fin.

        Supprime les paragraphes avant @debut_tpe@ (et le marqueur lui-même)
        du document original pour conserver tous les styles et la mise en forme.

        Args:
            docx_path: Chemin vers le fichier .docx du TRE.
            questions: Dictionnaire {clé: texte} des questions (non utilisé ici).

        Returns:
            Contenu du fichier .docx généré, sous forme de bytes.

        Raises:
            ValueError: Si le marqueur @debut_tpe@ n'est pas trouvé.
        """
        from copy import deepcopy
        from docx.oxml.ns import qn

        source_doc = DocxDocument(docx_path)
        debut_tpe_idx = self._find_debut_tpe_index(source_doc)

        if debut_tpe_idx is None:
            raise ValueError(
                "Le marqueur @debut_tpe@ n'a pas été trouvé dans le document."
            )

        # Supprimer les paragraphes avant @debut_tpe@ (inclus) du document original
        # On travaille directement sur le XML pour préserver les styles
        body = source_doc.element.body
        paragraphs_to_remove = []
        for idx in range(debut_tpe_idx + 1):  # +1 pour inclure @debut_tpe@ lui-même
            paragraphs_to_remove.append(source_doc.paragraphs[idx]._element)

        for p_element in paragraphs_to_remove:
            body.remove(p_element)

        # Sauvegarder en bytes
        buffer = io.BytesIO()
        source_doc.save(buffer)
        return buffer.getvalue()

    def extract_header(self, docx_path: str) -> bytes:
        """Extrait l'en-tête du TRE (avant @debut_tpe@).

        Crée un nouveau document contenant tous les paragraphes
        précédant le marqueur @debut_tpe@.

        Args:
            docx_path: Chemin vers le fichier .docx du TRE.

        Returns:
            Contenu du fichier .docx généré, sous forme de bytes.

        Raises:
            ValueError: Si le marqueur @debut_tpe@ n'est pas trouvé.
        """
        source_doc = DocxDocument(docx_path)
        debut_tpe_idx = self._find_debut_tpe_index(source_doc)

        if debut_tpe_idx is None:
            raise ValueError(
                "Le marqueur @debut_tpe@ n'a pas été trouvé dans le document."
            )

        # Créer un nouveau document
        new_doc = DocxDocument()

        # Copier les paragraphes avant @debut_tpe@
        for idx in range(debut_tpe_idx):
            source_para = source_doc.paragraphs[idx]
            self._copy_paragraph(source_para, new_doc)

        # Sauvegarder en bytes
        buffer = io.BytesIO()
        new_doc.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Méthodes privées
    # ------------------------------------------------------------------

    def _parse_annotation_key(self, key: str) -> tuple[str, str, bool]:
        """Parse la clé d'une annotation pour extraire type, suffix et is_custom.

        Règles :
        - Si la clé commence par '/', c'est une annotation personnalisée.
          Le type est la clé sans le '/', le suffix est vide.
        - Sinon, on sépare sur le premier '_' pour obtenir type et suffix.
          Ex: "dires_fratrie" → ("dires", "fratrie")
          Ex: "analyse" → ("analyse", "")

        Args:
            key: Clé brute de l'annotation (ex: "dires_fratrie", "/mon_annotation").

        Returns:
            Tuple (type, suffix, is_custom).
        """
        if key.startswith("/"):
            # Annotation personnalisée
            custom_name = key[1:]
            return custom_name, "", True

        # Annotation prédéfinie : séparer sur le premier '_'
        parts = key.split("_", 1)
        annotation_type = parts[0]

        if annotation_type in self.PREDEFINED_TYPES:
            suffix = parts[1] if len(parts) > 1 else ""
            return annotation_type, suffix, False

        # Type non reconnu : on le traite quand même, suffix complet
        suffix = parts[1] if len(parts) > 1 else ""
        return annotation_type, suffix, False

    def _find_debut_tpe_index(self, doc: DocxDocument) -> int | None:
        """Trouve l'index du paragraphe contenant @debut_tpe@.

        Args:
            doc: Document python-docx ouvert.

        Returns:
            Index du paragraphe ou None si non trouvé.
        """
        for idx, paragraph in enumerate(doc.paragraphs):
            if self.DEBUT_TPE_RE.match(paragraph.text):
                return idx
        return None

    def _copy_paragraph(self, source_para, target_doc: DocxDocument) -> None:
        """Copie un paragraphe source dans le document cible.

        Préserve le texte, le style et le formatage des runs.

        Args:
            source_para: Paragraphe source (python-docx Paragraph).
            target_doc: Document cible dans lequel ajouter le paragraphe.
        """
        new_para = target_doc.add_paragraph()

        # Copier le style du paragraphe
        if source_para.style:
            try:
                new_para.style = source_para.style.name
            except (KeyError, ValueError):
                # Le style n'existe pas dans le document cible
                pass

        # Copier l'alignement
        try:
            if source_para.alignment is not None:
                new_para.alignment = source_para.alignment
        except (ValueError, KeyError):
            # Certains alignements Word modernes ne sont pas supportés par python-docx
            pass

        # Copier les runs avec leur formatage
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            # Copier les propriétés de formatage
            if run.bold is not None:
                new_run.bold = run.bold
            if run.italic is not None:
                new_run.italic = run.italic
            if run.underline is not None:
                new_run.underline = run.underline
            if run.font.size is not None:
                new_run.font.size = run.font.size
            if run.font.name is not None:
                new_run.font.name = run.font.name
