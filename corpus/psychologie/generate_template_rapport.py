"""
Script de génération du template de rapport d'expertise psychologique judiciaire.
Exécuter ce script pour produire le fichier template_rapport_psychologie.docx

Usage : python generate_template_rapport.py

Prérequis : pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_template():
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Page de garde ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT D'EXPERTISE PSYCHOLOGIQUE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Expertise judiciaire ordonnée par")
    run.font.size = Pt(14)

    # Placeholder destinataire (extrait du QT)
    dest = doc.add_paragraph()
    dest.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dest.add_run("{{JURIDICTION_NOM}}")
    run.bold = True
    run.font.size = Pt(14)

    dest2 = doc.add_paragraph()
    dest2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dest2.add_run("{{JURIDICTION_ADRESSE}}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run("Référence dossier : {{REFERENCE_DOSSIER}}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Date du rapport : {{DATE_RAPPORT}}")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Informations de l'expert ---
    doc.add_heading("INFORMATIONS DE L'EXPERT", level=1)

    table_expert = doc.add_table(rows=5, cols=2)
    table_expert.style = 'Light Grid Accent 1'
    cells = table_expert.rows[0].cells
    cells[0].text = "Nom de l'expert"
    cells[1].text = "{{EXPERT_NOM}}"
    cells = table_expert.rows[1].cells
    cells[0].text = "Qualité"
    cells[1].text = "{{EXPERT_QUALITE}}"
    cells = table_expert.rows[2].cells
    cells[0].text = "N° inscription"
    cells[1].text = "{{EXPERT_NUMERO}}"
    cells = table_expert.rows[3].cells
    cells[0].text = "Cour d'appel"
    cells[1].text = "{{EXPERT_COUR_APPEL}}"
    cells = table_expert.rows[4].cells
    cells[0].text = "Coordonnées"
    cells[1].text = "{{EXPERT_COORDONNEES}}"

    doc.add_paragraph()

    # --- Destinataire ---
    doc.add_heading("DESTINATAIRE", level=1)

    table_dest = doc.add_table(rows=4, cols=2)
    table_dest.style = 'Light Grid Accent 1'
    cells = table_dest.rows[0].cells
    cells[0].text = "Juridiction"
    cells[1].text = "{{JURIDICTION_NOM}}"
    cells = table_dest.rows[1].cells
    cells[0].text = "Magistrat"
    cells[1].text = "{{MAGISTRAT_NOM}}"
    cells = table_dest.rows[2].cells
    cells[0].text = "Référence"
    cells[1].text = "{{REFERENCE_DOSSIER}}"
    cells = table_dest.rows[3].cells
    cells[0].text = "Date de la réquisition"
    cells[1].text = "{{DATE_REQUISITION}}"

    doc.add_paragraph()

    # --- Mention méthodologique ---
    doc.add_heading("MENTION MÉTHODOLOGIQUE", level=1)

    mention = doc.add_paragraph()
    mention.add_run(
        "Le présent rapport a été rédigé avec l'assistance d'un outil d'intelligence "
        "artificielle (Judi-Expert) utilisé comme aide à la structuration et à la rédaction, "
        "conformément à la méthodologie décrite dans le document de référence disponible sur "
        "le site central Judi-Expert. L'ensemble des analyses, conclusions et avis exprimés "
        "dans ce rapport relèvent de la seule responsabilité de l'expert signataire. "
        "L'outil IA a été utilisé pour : la génération du plan d'entretien, la mise en forme "
        "du rapport final et l'analyse critique des conclusions. Toutes les données d'expertise "
        "ont été traitées exclusivement sur le poste informatique de l'expert, sans transmission "
        "à des tiers."
    )

    doc.add_paragraph()

    # --- Objet de la mission ---
    doc.add_heading("1. OBJET DE LA MISSION", level=1)

    doc.add_paragraph(
        "Par réquisition en date du {{DATE_REQUISITION}}, {{MAGISTRAT_NOM}}, "
        "{{MAGISTRAT_QUALITE}} près {{JURIDICTION_NOM}}, a ordonné une expertise "
        "psychologique de :"
    )

    table_mec = doc.add_table(rows=3, cols=2)
    table_mec.style = 'Light Grid Accent 1'
    cells = table_mec.rows[0].cells
    cells[0].text = "Nom et prénom"
    cells[1].text = "{{MEC_NOM}} {{MEC_PRENOM}}"
    cells = table_mec.rows[1].cells
    cells[0].text = "Date de naissance"
    cells[1].text = "{{MEC_DATE_NAISSANCE}}"
    cells = table_mec.rows[2].cells
    cells[0].text = "Faits reprochés"
    cells[1].text = "{{FAITS_REPROCHES}}"

    doc.add_paragraph()

    # --- Questions du tribunal ---
    doc.add_heading("2. QUESTIONS DU TRIBUNAL", level=1)

    doc.add_paragraph(
        "Le tribunal a posé les questions suivantes auxquelles la présente expertise "
        "doit répondre :"
    )

    doc.add_paragraph("{{QUESTIONS_TRIBUNAL}}")

    doc.add_paragraph()

    # --- Déroulement de l'expertise ---
    doc.add_heading("3. DÉROULEMENT DE L'EXPERTISE", level=1)

    doc.add_paragraph("{{DEROULEMENT_EXPERTISE}}")

    doc.add_paragraph()

    # --- Anamnèse ---
    doc.add_heading("4. ANAMNÈSE", level=1)

    doc.add_heading("4.1 Histoire personnelle et familiale", level=2)
    doc.add_paragraph("{{ANAMNESE_PERSONNELLE}}")

    doc.add_heading("4.2 Parcours scolaire et professionnel", level=2)
    doc.add_paragraph("{{ANAMNESE_SCOLAIRE}}")

    doc.add_heading("4.3 Antécédents médicaux et psychologiques", level=2)
    doc.add_paragraph("{{ANAMNESE_MEDICALE}}")

    doc.add_heading("4.4 Antécédents judiciaires", level=2)
    doc.add_paragraph("{{ANAMNESE_JUDICIAIRE}}")

    doc.add_paragraph()

    # --- Examen clinique ---
    doc.add_heading("5. EXAMEN CLINIQUE", level=1)

    doc.add_heading("5.1 Présentation et contact", level=2)
    doc.add_paragraph("{{EXAMEN_PRESENTATION}}")

    doc.add_heading("5.2 Fonctionnement cognitif", level=2)
    doc.add_paragraph("{{EXAMEN_COGNITIF}}")

    doc.add_heading("5.3 Fonctionnement affectif et émotionnel", level=2)
    doc.add_paragraph("{{EXAMEN_AFFECTIF}}")

    doc.add_heading("5.4 Positionnement par rapport aux faits", level=2)
    doc.add_paragraph("{{EXAMEN_POSITIONNEMENT}}")

    doc.add_paragraph()

    # --- Tests psychométriques ---
    doc.add_heading("6. TESTS PSYCHOMÉTRIQUES (le cas échéant)", level=1)

    doc.add_paragraph("{{TESTS_PSYCHOMETRIQUES}}")

    doc.add_paragraph()

    # --- Analyse et discussion ---
    doc.add_heading("7. ANALYSE ET DISCUSSION", level=1)

    doc.add_paragraph("{{ANALYSE_DISCUSSION}}")

    doc.add_paragraph()

    # --- Réponses aux questions du tribunal ---
    doc.add_heading("8. RÉPONSES AUX QUESTIONS DU TRIBUNAL", level=1)

    doc.add_paragraph("{{REPONSES_QT}}")

    doc.add_paragraph()

    # --- Conclusion ---
    doc.add_heading("9. CONCLUSION", level=1)

    doc.add_paragraph("{{CONCLUSION}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Signature ---
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = signature.add_run("Fait à {{LIEU}}, le {{DATE_RAPPORT}}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    sig_name = doc.add_paragraph()
    sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig_name.add_run("{{EXPERT_NOM}}")
    run.bold = True
    run.font.size = Pt(12)

    sig_qual = doc.add_paragraph()
    sig_qual.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig_qual.add_run("{{EXPERT_QUALITE}}")
    run.font.size = Pt(12)

    # --- Sauvegarde ---
    output_path = "corpus/psychologie/template_rapport_psychologie.docx"
    doc.save(output_path)
    print(f"Template généré : {output_path}")


if __name__ == "__main__":
    create_template()
